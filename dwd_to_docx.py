#!/usr/bin/env python3
"""
DWD to DOCX Converter  –  DavkaWriter .dwd → Microsoft Word .docx

Usage:
    python dwd_to_docx.py input.dwd [output.docx] [--no-nikud] [--no-trup]

    Nikud and trup are included by default. Use --no-nikud / --no-trup to omit.

DavkaWriter format versions (auto-detected by run signature):
    Format A  49 80 01 00 01 00 00 00 02  — DavkaWriter for Windows (original)
    Format B  08 82 01 00 01 00 00 00 02  — DavkaWriter Gold / newer editions
    Format C  11 81 01 00 01 00 00 00 02  — DavkaWriter alternate edition

All formats share the same Hebrew consonant encoding (DAVKA_MAP).
Hebrew vs English detection: for Format A, the style byte determines the font.
For Formats B/C, content-based detection is used (alef byte 0x60 = definitive Hebrew marker).

Nikud bytes (follow the consonant they vowelize):
    0x9f=holam  0xa1=hataf-segol  0xa2=hataf-kamatz  0xa3=hataf-patach
    0xa4=segol  0xa5=tsere  0xa6=hiriq  0xa7=sheva
    0xa8=kamatz  0xa9=patach  0xaa=kubutz
    (confirmed by byte-alignment vs vocalized Mishna Shabbos 1:1)

Trup bytes (ta'amei hamikra / cantillation marks):
    0x84=QARNEY PARA  0x85=MERKHA KEFULA  0xab=SOF PASUK  0xac=MERKHA
    0xad=TIPEHA       0xae=ETNAHTA        0xaf=DARGA       0xb2=TEVIR
    0xb3=MAHAPAKH     0xb4=MUNAH          0xb5=YETIV       0xb6=YERACH BEN YOMO
    0xb8=QADMA        0xb9=GERESH         0xba=GERSHAYIM   0xbb=ZARQA
    0xbc=SEGOL(trup)  0xbd=ZAQEF QATAN   0xbe=ZAQEF GADOL  0xbf=PAZER
    0xc0=REVIA        0xc1=TELISHA QETANA  0xc2=TELISHA GEDOLA  0xc3=SHALSHELET
    0xc6=PASHTA       0xce=MAQAF
    (confirmed by Sha'ar Hata'amim worksheets, 21368)
"""

import io, re, struct, sys, zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Consonant map ─────────────────────────────────────────────────────────────
DAVKA_MAP = {
    0x60:'א', 0x61:'ב', 0x62:'ג', 0x63:'ד', 0x64:'ה', 0x65:'ו', 0x66:'ז',
    0x67:'ח', 0x68:'ט', 0x69:'י', 0x6a:'כ', 0x6b:'כ', 0x6c:'ל', 0x6d:'ם',
    0x6e:'מ', 0x6f:'ן', 0x70:'נ', 0x71:'ס', 0x72:'ע', 0x73:'ף', 0x74:'פ',
    0x75:'ץ', 0x76:'צ', 0x77:'ק', 0x78:'ר', 0x79:'ש', 0x7a:'ת',
    # uppercase dagesh variants (confirmed from text)
    0x41:'ב\u05BC', 0x49:'י\u05BC', 0x4d:'כ\u05BC', 0x4e:'ל\u05BC',
    0x4f:'מ\u05BC', 0x51:'ס\u05BC', 0x52:'פ\u05BC',
    0x57:'ש\u05C1',   # shin-dot
    # uppercase alternates
    0x42:'ג', 0x43:'ד', 0x44:'ה', 0x45:'ו', 0x46:'ו', 0x47:'ג', 0x48:'ט',
    0x4a:'כ', 0x4b:'ך', 0x4c:'ל', 0x50:'נ', 0x53:'ש', 0x54:'ק', 0x55:'ש',
    0x56:'צ', 0x58:'ש', 0x59:'ת', 0x5a:'פ',
}

NIKUD_MAP = {
    0x9f:'\u05B9', 0xa1:'\u05B1', 0xa2:'\u05B3', 0xa3:'\u05B2',
    0xa4:'\u05B6', 0xa5:'\u05B5', 0xa6:'\u05B4', 0xa7:'\u05B0',
    0xa8:'\u05B8', 0xa9:'\u05B7', 0xaa:'\u05BB', 0xfb:'\u05B9',
}

TRUP_MAP = {
    0x84:'\u059F', 0x85:'\u05A6', 0xab:'\u05C3', 0xac:'\u05A5',
    0xad:'\u0596', 0xae:'\u0591', 0xaf:'\u05A7', 0xb2:'\u059B',
    0xb3:'\u05A4', 0xb4:'\u05A3', 0xb5:'\u059A', 0xb6:'\u05AA',
    0xb8:'\u05A8', 0xb9:'\u059C', 0xba:'\u059E', 0xbb:'\u0598',
    0xbc:'\u0592', 0xbd:'\u0594', 0xbe:'\u0595', 0xbf:'\u05A1',
    0xc0:'\u0597', 0xc1:'\u05A9', 0xc2:'\u05A0', 0xc3:'\u0593',
    0xc6:'\u0599', 0xce:'\u05BE',
}

# Style constants (Format A only — Formats B/C use content detection)
HEB_STYLES      = {0x20,0x23,0x25,0x28,0x29,0x2b,0x2c}
KEYWORD_HEB_STY = {0x23,0x25}
KEYWORD_DEF_STY = 0x24
HEB_HEADING_STY = 0x2b
MISHNA_STYS     = {0x29,0x20}
SECTION_HDR_STY = 0x27
BOLD_STY        = 0x1f

# ── Format detection ──────────────────────────────────────────────────────────
# Known DavkaWriter run-signature variants (used as a lookup for naming).
# The parser also auto-detects unknown variants dynamically.
_KNOWN_FORMATS = {
    bytes([0x49,0x80,0x01,0x00,0x01,0x00,0x00,0x00,0x02]): 'Format A (DavkaWriter original, 49-80)',
    bytes([0x08,0x82,0x01,0x00,0x01,0x00,0x00,0x00,0x02]): 'Format B (DavkaWriter Gold, 08-82)',
    bytes([0x11,0x81,0x01,0x00,0x01,0x00,0x00,0x00,0x02]): 'Format C (DavkaWriter alternate, 11-81)',
    bytes([0x49,0x81,0x01,0x00,0x01,0x00,0x00,0x00,0x02]): 'Format D (DavkaWriter poster/style, 49-81)',
    bytes([0x25,0x82,0x01,0x00,0x01,0x00,0x00,0x00,0x02]): 'Format E (DavkaWriter 25-82)',
}

# The invariant suffix shared by ALL DavkaWriter run signatures.
# Bytes [2..8] of the 9-byte header are always: 01 00 01 00 00 00 02
_RUN_SIG_SUFFIX = bytes([0x01,0x00,0x01,0x00,0x00,0x00,0x02])

def _detect_format(data):
    """Auto-detect DWD run/para signatures from the file itself.

    Every DWD run signature is 9 bytes: [b0][b1] 01 00 01 00 00 00 02
    where b0/b1 identify the format variant. Rather than matching a fixed
    list, we scan for any 9-byte pattern ending in the invariant suffix
    and pick the one with the highest count — that's the run signature.

    The para signature is always [b0-2][b1] 01 00 01 (5 bytes), sharing
    b1 and using b0-2 as the first byte (observed across all known formats:
    49→47, 08→06, 11→0f, 49→47, 25→23).

    Returns (run_sig, para_sig, fmt_name, use_style_detection).
    """
    # Count all candidate 9-byte patterns
    counts = {}
    pos = 0
    suffix = _RUN_SIG_SUFFIX
    slen = len(suffix)
    while True:
        p = data.find(suffix, pos)
        if p < 0: break
        if p >= 2:
            sig = data[p-2:p+slen]   # 2 prefix bytes + 7 suffix bytes = 9 total
            if len(sig) == 9:
                counts[sig] = counts.get(sig, 0) + 1
        pos = p + 1

    if not counts:
        return None, None, 'Unknown (no run signatures found)', False

    # Pick the most frequent signature
    run_sig = max(counts, key=counts.__getitem__)

    # Derive the para signature: b1 is shared; b0_para = b0_run - 2 (observed pattern)
    b0_run, b1_run = run_sig[0], run_sig[1]
    b0_para = b0_run - 2
    para_sig = bytes([b0_para, b1_run, 0x01, 0x00, 0x01])

    # Name it
    if run_sig in _KNOWN_FORMATS:
        fmt_name = _KNOWN_FORMATS[run_sig]
    else:
        fmt_name = f'Unknown variant ({b0_run:02x}-{b1_run:02x}, {counts[run_sig]} occurrences)'

    # Only Format A uses style-byte language detection
    use_style_detection = (run_sig == bytes([0x49,0x80,0x01,0x00,0x01,0x00,0x00,0x00,0x02]))

    return run_sig, para_sig, fmt_name, use_style_detection

_BAD_XML = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]')

# ── Style table parser ────────────────────────────────────────────────────────
# Maps Davka font names to their best Word-compatible equivalents.
_FONT_SUBS = {
    'Davka Stam':       'Ezra SIL',
    'Davka FrankRuhl':  'Frank Ruehl CLM',
    'Davka Hadasah':    'Noto Serif Hebrew',
    'Davka David':      'David',
    'Davka Rashi':      'Noto Rashi Hebrew',
    'Davka Kastel':     'Noto Serif Hebrew',
    'Davka Yerushalmy': 'Noto Serif Hebrew',
    'Davka Meir':       'Noto Sans Hebrew',
    'Davka Siddur':     'Ezra SIL',
    'Arial':            'Arial',
    'Times New Roman':  'Times New Roman',
    'Calibri':          'Calibri',
    'David':            'David',
    'Courier New':      'Courier New',
    'Courier':          'Courier New',
}
_FONT_IS_HEBREW = {k for k in _FONT_SUBS if k.startswith('Davka') or k == 'David'}

_CDOC_FONT_RE = re.compile(
    b'(Davka [\x20-\x7e]+|Arial|Times New Roman|Calibri|David|Courier New|Courier)\x00'
)

def extract_header_footer_events(data, run_sig, para_sig):
    """Extract runs/paras from CHeader and CFooter sections.

    Returns a dict {'header': [events], 'footer': [events]} where events
    are in the same format as parse_dwd returns.

    Scanning stops at structural markers like CFootnoteSet or CDocRuler,
    but skips over inline field markers like CFieldPageNo so the full
    header/footer content is captured.

    DavkaWriter stores per-page headers sequentially in a single CHeader
    block; Word only supports one header per section.  To avoid putting
    page-specific content into a document-wide header, we stop at the
    first "double PARA" break, which usually separates distinct headers.
    """
    SKIP_MARKERS = {
        b'CFieldPageNo', b'CFieldGraphic', b'CFieldTextBox',
        b'CFieldTextLink', b'CFieldAutoShape', b'CFieldCurvyText',
    }
    END_MARKERS = {
        b'CFootnoteSet', b'CDocRuler', b'CPageObject', b'CDocStyle',
        b'CParagList', b'CHeader', b'CFooter', b'CTableRow',
        b'CParagObject', b'CStyleRun',
    }

    result = {}
    for label, marker in [('header', b'CHeader'), ('footer', b'CFooter')]:
        pos = data.find(marker)
        if pos < 0:
            continue
        start = pos + len(marker)
        end_scan = min(start + 3000, len(data))
        events = []
        i = start
        # Stop after first double PARA (for headers with per-page content)
        max_content_runs = 5  # cap total content runs
        consecutive_paras = 0
        while i < end_scan - 15:
            # MFC class marker
            if data[i:i+4] == b'\xff\xff\x00\x00':
                nlen = data[i+4]
                if 0 < nlen < 30:
                    nm = data[i+6:i+6+nlen]
                    if (nm.startswith(b'C') and nm != marker
                        and all(32 <= b < 127 for b in nm)):
                        if nm in END_MARKERS:
                            break
                        if nm in SKIP_MARKERS:
                            i += 6 + nlen
                            continue
                        break
            # Run signature
            if data[i:i+9] == run_sig:
                sty, hi, ln = data[i+9], data[i+10], data[i+11]
                if hi in (0, 1) and ln > 0 and i + 12 + ln <= len(data):
                    events.append({
                        'type': 'run', 'style': sty,
                        'bytes': data[i+12:i+12+ln],
                        'use_style': False,
                        'hi_flag': hi,
                    })
                    consecutive_paras = 0
                    # Stop if we have too many runs (likely picking up body)
                    content_runs = sum(1 for ev in events if ev['type'] == 'run')
                    if content_runs >= max_content_runs:
                        break
                    i += 12 + ln
                    continue
            # Para signature
            if data[i:i+5] == para_sig:
                events.append({'type': 'para'})
                consecutive_paras += 1
                # Two consecutive paras = end of first header/footer block
                if consecutive_paras >= 2 and events and any(
                    ev['type'] == 'run' for ev in events
                ):
                    break
                i += 5
                continue
            i += 1
        if events:
            result[label] = events
    return result


def parse_style_table(data):
    """Parse the CDocStyle section and return a dict: style_index → props dict.

    Each props dict has:
      font      – Word-compatible font name (string)
      size_pt   – font size in points (float), 0 if unknown
      bold      – bool
      italic    – bool
      underline – bool
      is_hebrew – bool (True for Davka Hebrew fonts)

    CDocStyle record layout (24 bytes of header before null-terminated font name):
      [0]     separator byte from previous record
      [1]     flags_hi  — bit 7 (0x80) = bold
      [2]     flags_lo  — bit 6 (0x40) = italic, bit 4 (0x10) = underline
      [3-5]   ???
      [6]     font_id (informational; we use the font name instead)
      [7]     0x00
      [8-10]  RGB (always 0xffffff in files seen so far — not decoded)
      [11]    0x00
      [12-19] reserved
      [20-23] LE32 size in tenths-of-a-point (e.g. 1200 = 12.0pt)
      [24+]   font name, null-terminated
    """
    cdoc = data.find(b'CDocStyle')
    if cdoc < 0:
        return {}

    styles = {}
    for m in _CDOC_FONT_RE.finditer(data, cdoc, cdoc + 200_000):
        fpos  = m.start()
        fname = m.group(1).decode('ascii', 'replace').rstrip()

        hdr_start = fpos - 24
        if hdr_start < cdoc:
            continue
        hdr = data[hdr_start:fpos]
        if len(hdr) < 24:
            continue

        flags_hi  = hdr[1]
        flags_lo  = hdr[2]
        size_raw  = struct.unpack_from('<I', hdr, 20)[0]
        size_pt   = round(size_raw / 100.0, 1) if 50 <= size_raw <= 10000 else 0

        bold      = bool(flags_hi & 0x80)
        italic    = bool(flags_lo & 0x40)
        underline = bool(flags_lo & 0x10)
        is_heb    = fname in _FONT_IS_HEBREW
        word_font = _FONT_SUBS.get(fname, fname)

        idx = len(styles)
        styles[idx] = {
            'font': word_font,
            'size_pt': size_pt,
            'bold': bold,
            'italic': italic,
            'underline': underline,
            'is_hebrew': is_heb,
        }

    return styles


def _normalize_style_sizes(style_table, events):
    """Normalize font sizes based on usage frequency.

    DavkaWriter sometimes records very large nominal sizes (e.g. 76pt) for
    body text styles that render at smaller sizes in the actual document.
    Heavily-used styles are body text and need to be capped at body sizes
    to produce readable Word output.

    Heuristic:
      - Most-used styles (>200 runs, >2000 bytes) → BODY: 16pt Heb, 12pt Eng
      - Medium usage (>50 runs) → MID: 22pt Heb, 18pt Eng
      - Low usage → DISPLAY: keep up to 36pt
    """
    from collections import Counter

    sty_runs = Counter()
    sty_bytes = Counter()
    for ev in events:
        if ev.get('type') == 'run':
            sty = ev['style']
            sty_runs[sty] += 1
            sty_bytes[sty] += len(ev.get('bytes', b''))

    normalized = {}
    for sty, props in style_table.items():
        new_props = dict(props)
        size = props.get('size_pt', 0)
        is_heb = props.get('is_hebrew', False)
        runs = sty_runs.get(sty, 0)
        bytes_ = sty_bytes.get(sty, 0)

        if runs > 200 and bytes_ > 2000:
            # Body text: cap aggressively
            cap = 16.0 if is_heb else 12.0
            if size > cap:
                new_props['size_pt'] = cap
        elif runs > 50:
            # Mid-level
            cap = 22.0 if is_heb else 18.0
            if size > cap:
                new_props['size_pt'] = cap
        else:
            # Display/poster: cap at 36pt
            if size > 36.0:
                new_props['size_pt'] = 36.0

        normalized[sty] = new_props

    return normalized

# ── Fix 1: Windows-1252 special characters ────────────────────────────────────
# These bytes appear in English runs and must be mapped to proper Unicode.
_WIN1252 = {
    0x80: '\u20AC',  # €
    0x85: '\u2026',  # …
    0x91: '\u2018',  # '  left single quotation mark
    0x92: '\u2019',  # '  right single quotation mark (apostrophe)
    0x93: '\u201C',  # "  left double quotation mark
    0x94: '\u201D',  # "  right double quotation mark
    0x96: '\u2013',  # –  en dash
    0x97: '\u2014',  # —  em dash
    0xA0: '\u00A0',  # non-breaking space
    0xA9: '\u00A9',  # ©  copyright
    0xAE: '\u00AE',  # ®  registered
    0xB0: '\u00B0',  # °  degree
    0xB1: '\u00B1',  # ±  plus-minus
    0xBC: '\u00BC',  # ¼
    0xBD: '\u00BD',  # ½
    0xBE: '\u00BE',  # ¾
    0x99: '\u2122',  # ™  trademark
}

# ── Fix 2: Davka formatting-marker detection ──────────────────────────────────
# DavkaWriter stores inline style/color instructions as short literal text runs.
# These are never real content and must be stripped from output.
_DAVKA_MARKERS = frozenset({
    'low', 'Low', 'low ', 'Low ', 'high', 'High', 'high ', 'High ',
    'lowlow', 'HighHigh',
    # Hebrew-decoded internal codes (confirmed across multiple files)
    'לןק', 'טיחט', 'לןק ', 'טיחט ',
    # All-caps ASCII codes
    'TTTO',
})

def _is_marker_run(text):
    """Return True if this decoded text is a Davka internal formatting marker."""
    t = text.strip()
    if not t:
        return False
    # Exact known markers
    if t in _DAVKA_MARKERS:
        return True
    # Pattern: 'low N' or 'high N' or 'לןק Nr' or 'טיחט Nr' (music/style markers)
    for prefix in ('low ', 'Low ', 'high ', 'High ', 'לןק ', 'טיחט '):
        if t.startswith(prefix) and len(t) <= len(prefix) + 4:
            return True
    return False


# ── Parser ────────────────────────────────────────────────────────────────────
def _is_hebrew_content(raw):
    """Content heuristic for runs where style classification is ambiguous.

    Strong Hebrew signals:
      - any nikud or trup byte adjacent to Hebrew consonant bytes → Hebrew
      - isolated alef byte (0x60) surrounded by Hebrew bytes → Hebrew
    Strong English signals:
      - starts with digit → English (lists)
      - all-uppercase first word that forms an English word → English
    Default: Hebrew (safer for this corpus, since Davka uses ASCII bytes
    0x60-0x7a AND uppercase 0x40-0x5a for Hebrew letters).
    """
    if not raw: return False

    DIACRIT = set(NIKUD_MAP) | set(TRUP_MAP)
    HEB_CONSONANTS = set(range(0x40, 0x5B)) | set(range(0x60, 0x7B))

    # Strong English signal: starts with digit
    if 0x30 <= raw[0] <= 0x39:
        return False

    # Nikud/trup adjacent to Hebrew consonant = definitively Hebrew
    has_diacrit = any(b in DIACRIT for b in raw)
    has_heb_consonant = any(b in HEB_CONSONANTS for b in raw)

    if has_diacrit and has_heb_consonant:
        # Both diacritic and Hebrew letter present → likely Hebrew
        # But if ASCII text dominates with explicit English words, it's English
        text = raw.decode('latin-1', 'replace')
        eng_words = (' the ', ' and ', ' from ', ' for ', ' with ', ' that ',
                     ' you ', ' is ', ' was ', ' to ', ' of ', ' by ', ' did ')
        padded = ' ' + text.lower() + ' '
        if any(w in padded for w in eng_words):
            return False
        return True

    # Strong English signal: high ratio of "English-only" chars
    all_low_ascii = all(b < 0x80 for b in raw)
    if all_low_ascii:
        text = bytes(raw).decode('ascii', 'replace').lower()
        english_words = (' the ', ' and ', ' from ', ' for ', ' with ', ' that ',
                         ' you ', ' your ', ' is ', ' was ', ' to ', ' of ',
                         ' in ', ' it ', ' a ', ' on ', ' be ', ' will ',
                         ' are ', ' this ', ' an ', ' by ', ' he ', ' they ',
                         ' them ', ' her ', ' his ', ' who ', ' what ',
                         ' which ', ' did ')
        padded = ' ' + text + ' '
        if any(w in padded for w in english_words):
            return False
        # Note: we deliberately do NOT use a "lowercase-start ratio" rule here —
        # Davka Hebrew encoding uses bytes 0x60–0x7a which ARE ASCII lowercase,
        # so looks-like-lowercase-words is not sufficient evidence of English.

    # Alef byte in Hebrew context
    if 0x60 in raw:
        return True

    return True   # default Hebrew


def _build_style_map(events):
    """Two-pass style→language map from unambiguous runs in this file.
    A style confirmed Hebrew anywhere is treated as Hebrew everywhere.
    Uses specific nikud/trup byte-sets rather than generic high-byte check
    to avoid mis-classifying English runs with Windows special characters.
    """
    DIACRIT = set(NIKUD_MAP) | set(TRUP_MAP)
    style_heb = {}
    for ev in events:
        if ev['type'] != 'run': continue
        raw, sty = ev['bytes'], ev['style']
        if any(b in DIACRIT for b in raw) or 0x60 in raw:
            style_heb[sty] = True                  # definitively Hebrew
        elif 0x41 <= raw[0] <= 0x5A and all(0x20 <= b <= 0x7E for b in raw):
            if sty not in style_heb:
                style_heb[sty] = False             # English unless contradicted
    return style_heb


def parse_dwd(data):
    """Parse DWD binary into (events, format_name).

    Auto-detects the DavkaWriter format variant from the file itself by
    scanning for the invariant 7-byte run-signature suffix shared by all
    known variants. Works on both known and previously-unseen format bytes.
    """
    run_sig, para_sig, fmt_name, use_style_detection = _detect_format(data)

    if run_sig is None:
        # No recognisable structure — return empty event list
        return [], fmt_name

    # JPEG signatures to detect embedded images
    JPEG_SIGS = (b'\xFF\xD8\xFF\xE0', b'\xFF\xD8\xFF\xE1', b'\xFF\xD8\xFF\xDB')

    events, i, n = [], 0, len(data)
    while i < n - 12:
        # ── Image detection ───────────────────────────────────────────────────
        if data[i:i+3] == b'\xFF\xD8\xFF' and i + 4 < n and data[i:i+4] in JPEG_SIGS:
            eoi = data.find(b'\xFF\xD9', i + 4)
            if eoi > 0 and eoi - i > 5000:   # skip tiny false positives
                raw = data[i:eoi+2]
                try:
                    from PIL import Image as _PIL
                    img = _PIL.open(io.BytesIO(raw))
                    w, h = img.size
                    if w > 50 and h > 50:
                        events.append({'type': 'image', 'raw': raw,
                                       'width': w, 'height': h, 'fmt': 'JPEG'})
                        i = eoi + 2; continue
                except Exception:
                    pass
        # ── Run detection ─────────────────────────────────────────────────────
        # Run signature layout (9 bytes) + [sty:1][hi:1][ln:1][content:ln bytes]
        # hi == 0: "normal" run (body text)
        # hi == 1: "styled" run, often Hebrew content with its own style index
        # (Format E files use both hi=0 and hi=1 extensively; hi=1 runs were
        # ignored in earlier versions and caused significant content loss.)
        if data[i:i+9] == run_sig:
            sty, hi, ln = data[i+9], data[i+10], data[i+11]
            if hi in (0, 1) and ln > 0 and i+12+ln <= n:
                events.append({
                    'type': 'run', 'style': sty,
                    'bytes': data[i+12:i+12+ln],
                    'use_style': use_style_detection,
                    'hi_flag': hi,
                })
                i += 12 + ln; continue
        if data[i:i+5] == para_sig:
            events.append({'type': 'para'})
            i += 5; continue
        i += 1

    # For non-Format-A: second pass to tag each event with is_hebrew
    if not use_style_detection:
        style_map = _build_style_map(events)
        for ev in events:
            if ev['type'] == 'run':
                raw, sty = ev['bytes'], ev['style']
                DIACRIT = set(NIKUD_MAP) | set(TRUP_MAP)
                if any(b in DIACRIT for b in raw) or 0x60 in raw:
                    ev['is_hebrew'] = True
                elif sty in style_map:
                    ev['is_hebrew'] = style_map[sty]
                else:
                    ev['is_hebrew'] = _is_hebrew_content(raw)

    return events, fmt_name


# ── Decoder ───────────────────────────────────────────────────────────────────
def _clean(s): return _BAD_XML.sub('', s)

def decode_heb(raw, with_nikud=True, with_trup=True):
    out = []
    for b in raw:
        if b in DAVKA_MAP:
            out.append(DAVKA_MAP[b])
        elif b == 0x20:
            out.append(' ')
        elif b < 0x80 and chr(b) in ',.;:!?\'"()-/0123456789\'':
            out.append(chr(b))
        elif b in NIKUD_MAP and with_nikud:
            # Attach to previous letter if any; otherwise output on a base char
            # so trup posters can show standalone marks visibly.
            if out and out[-1] != ' ':
                out.append(NIKUD_MAP[b])
            else:
                # Use combining-char dotted circle base (Unicode 0x25CC) to make
                # standalone marks visible in Word
                out.append('\u25CC' + NIKUD_MAP[b])
        elif b in TRUP_MAP and with_trup:
            if out and out[-1] != ' ':
                out.append(TRUP_MAP[b])
            else:
                out.append('\u25CC' + TRUP_MAP[b])
    return _clean(''.join(out))

def _decode_ascii(raw):
    """Decode an English/ASCII run, mapping Windows-1252 special chars correctly."""
    out = []
    for b in raw:
        if b in _WIN1252:
            out.append(_WIN1252[b])
        elif 0x20 <= b <= 0x7E:
            out.append(chr(b))
        elif b in (0x09, 0x0A, 0x0D):
            out.append(chr(b))
        # skip other control chars
    return _clean(''.join(out))

def decode_run(ev, with_nikud=True, with_trup=True):
    sty, raw = ev['style'], ev['bytes']
    start = 0
    if raw and raw[0] < 0x20:
        start = 1
        if len(raw) > 1 and raw[1] < 0x20:
            start = 2
    content = raw[start:]

    # Strong Hebrew signal: presence of Davka NIKUD or TRUP bytes (0x9f-0xce range).
    # If found, this is definitely Davka Hebrew regardless of what the style says.
    has_davka_diacritic = any(b in NIKUD_MAP or b in TRUP_MAP for b in content)

    if ev.get('use_style', False):
        if sty in HEB_STYLES or has_davka_diacritic:
            return decode_heb(content, with_nikud, with_trup)
        return _decode_ascii(raw)

    # Formats B–E: use pre-computed is_hebrew tag if available
    is_heb = ev.get('is_hebrew', _is_hebrew_content(content))

    # Strong override: Davka diacritic bytes always mean Hebrew
    if has_davka_diacritic:
        is_heb = True

    # Per-run override: even if the style is classified Hebrew, some runs contain
    # clean English text (English typed in a Hebrew-font style).  Detect by
    # checking whether content would decode as readable English ASCII.
    if is_heb and not has_davka_diacritic and _is_clean_english(content):
        is_heb = False

    # Short ASCII content (1-3 chars) consisting only of Davka-uppercase
    # letter codes (R='פּ', W='שׁ', S='שׁ', K='ך', E='ו', m='ם') is Hebrew with
    # dagesh, NOT English.  Override even if is_hebrew=False was set.
    if not is_heb and 1 <= len(content) <= 3:
        # Davka uppercase Hebrew letters
        DAVKA_UPPER = set(b'RWSKEm')
        # Allowed extra chars: space and possibly ASCII letter for compound
        if all(b in DAVKA_UPPER or b == 0x20 for b in content):
            # All bytes are Davka uppercase or space → treat as Hebrew
            is_heb = True

    if is_heb:
        return decode_heb(content, with_nikud, with_trup)
    return _decode_ascii(content)


def _is_clean_english(raw):
    """Conservative heuristic: does this look like clean English ASCII text?

    Used to override a style's Hebrew classification when a run actually
    contains English content.  Must be CONSERVATIVE — Davka Hebrew encoding
    uses bytes 0x60–0x7a which ARE ASCII lowercase, so "looks like letters"
    is not enough.  We require explicit English-word evidence.
    """
    if not raw: return False

    # English with leading Win1252 special char (©, ®, °, ™, etc.)
    SPECIAL_AT_START = {0xa9, 0xae, 0xb0, 0x99, 0xa2, 0xa3, 0xa5}
    if raw[0] in SPECIAL_AT_START and len(raw) >= 4:
        rest = raw[1:]
        ascii_count = sum(1 for b in rest if 0x20 <= b < 0x80)
        if ascii_count > len(rest) * 0.7:
            return True

    # If the bytes look like UTF-8 encoded English (with smart quotes etc.),
    # decode as UTF-8 and check if remaining content is ASCII-dominant English.
    # Common English Unicode punct: smart quotes, em/en-dash, ellipsis, ©, ®, °, ™
    ENGLISH_UNICODE = (
        '\u2018\u2019\u201c\u201d'  # smart quotes
        '\u2013\u2014'              # em/en dashes
        '\u2026'                    # ellipsis
        '\u00a0'                    # non-breaking space
        '\u00a9\u00ae\u00b0\u2122'  # ©, ®, °, ™
        '\u00a2\u00a3\u00a5'        # ¢, £, ¥
        '\u00bc\u00bd\u00be'        # ¼, ½, ¾
    )
    # Try UTF-8 first, then Win1252 (DavkaWriter uses Win1252 high bytes for
    # smart quotes and other typography in English text).
    decoded_text = None
    for encoding in ('utf-8', 'windows-1252'):
        try:
            t = raw.decode(encoding)
            if encoding == 'windows-1252':
                # Win1252 decodes everything; only accept if NO Hebrew-range
                # high bytes are present (those are real Davka Hebrew encoding)
                # AND content has typical English punctuation Unicode points
                has_eng_punct = any(c in ENGLISH_UNICODE for c in t)
                if not has_eng_punct:
                    continue
            decoded_text = t
            break
        except UnicodeDecodeError:
            continue

    if decoded_text and all(c < '\u0080' or c in ENGLISH_UNICODE for c in decoded_text):
        # All chars are ASCII or English-only Unicode → treat as ASCII
        stripped = ''.join(c if c < '\u0080' else "'" for c in decoded_text)
        raw = stripped.encode('ascii', 'ignore')
        # Check leading-special after Unicode normalization
        if decoded_text and decoded_text[0] in '\u00a9\u00ae\u00b0\u2122' and len(decoded_text) >= 3:
            rest_chars = decoded_text[1:]
            ascii_count = sum(1 for c in rest_chars if c < '\u0080')
            if ascii_count > len(rest_chars) * 0.7:
                return True

    if not all(b < 0x80 for b in raw): return False
    if b'\x60' in raw:  # alef byte — strong Hebrew indicator
        return False
    text = bytes(raw).decode('ascii', 'replace')
    letter_count = sum(1 for c in text if c.isalpha())
    if letter_count < 1: return False

    # Words with internal apostrophe (e.g. "Par'oh", "C'na'an", "we'll", "doesn't")
    # are almost certainly English/transliteration
    if "'" in text and any(c.isalpha() for c in text):
        # Check that apostrophe is between letters
        for i in range(1, len(text) - 1):
            if text[i] == "'" and text[i-1].isalpha() and text[i+1].isalpha():
                return True

    # Single uppercase letter or tiny run starting with uppercase — likely English
    stripped = text.strip()
    if stripped in ('I', 'A', 'He', 'She', 'It', 'We', 'My', 'Me', 'Us',
                    'My ', 'We ', 'I ', 'A '):
        return True

    lower = text.lower()
    padded = ' ' + lower + ' '

    eng_words = (' the ', ' and ', ' from ', ' for ', ' with ', ' that ',
                 ' you ', ' your ', ' is ', ' was ', ' to ', ' of ',
                 ' in ', ' it ', ' on ', ' be ', ' will ', ' are ',
                 ' this ', ' an ', ' by ', ' he ', ' they ', ' them ',
                 ' her ', ' his ', ' who ', ' what ', ' which ',
                 ' did ', ' have ', ' has ', ' had ', ' can ', ' not ',
                 ' but ', ' so ', ' as ', ' at ', ' one ', ' two ',
                 ' three ', ' four ', ' five ', ' no ', ' yes ',
                 ' my ', ' me ', ' we ', ' us ', ' or ', ' all ',
                 ' out ', ' up ', ' go ', ' do ', ' if ',
                 ' hehim', ' heme', ' heyou', ' himher', ' theyhim',
                 ' themhis', ' hehis', ' ithim',
                 ' sokol ', ' inc ', ' ltd ', ' corp ', ' davka ', ' co ')
    if any(w in padded for w in eng_words):
        return True

    if len(text) >= 2 and text[0].isdigit() and text[1] in ').':
        return True

    # Title Case English
    clean_text = text.rstrip('\x00\x01\x02\x03\x04\x05')
    clean_text = clean_text.lstrip('\x00\x01\x02\x03\x04\x05')
    words = clean_text.split()
    if len(words) >= 2:
        cap_words = sum(1 for w in words if w and 'A' <= w[0] <= 'Z')
        if cap_words >= 2 and cap_words >= len(words) * 0.5:
            valid = True
            for w in words:
                if w and 'A' <= w[0] <= 'Z':
                    rest = w[1:]
                    if rest and not all(
                        ('a' <= c <= 'z') or c in '.,;:-\'"!?()' for c in rest
                    ):
                        valid = False
                        break
            if valid:
                return True

    # Single English word(s)
    clean2 = text.rstrip('\x00\x01\x02\x03\x04\x05')
    clean2 = clean2.lstrip('\x00\x01\x02\x03\x04\x05')
    text_lower = clean2.lower().strip(' .,;:!?\'"()')

    # For multi-word phrases, check each individual word
    is_single_word = len(text_lower.split()) == 1

    vowel_patterns = ('oo', 'ee', 'ai', 'ea', 'ou', 'ie', 'oa', 'ay', 'ey',
                      'oi', 'ow', 'igh', 'ough')
    suffixes = ('ed', 'ing', 'er', 'ly', 'tion', 'sion', 'ness',
                'ment', 'able', 'ful', 'ous', 'ize', 'ity')
    common_short = {'is', 'it', 'at', 'on', 'in', 'an', 'as', 'by', 'be',
                    'we', 'he', 'me', 'my', 'so', 'if', 'or', 'no', 'go',
                    'do', 'us', 'i', 'a', 'the', 'and', 'for', 'you',
                    'are', 'was', 'his', 'her', 'has', 'had', 'not',
                    'but', 'can', 'all', 'one', 'two', 'said', 'who',
                    'now', 'new', 'old', 'top', 'big', 'six', 'too',
                    'put', 'end', 'get', 'use', 'man', 'way', 'how',
                    'why', 'let', 'set', 'yes', 'she', 'them', 'him',
                    'our', 'their', 'than', 'then', 'made', 'make',
                    'were', 'will', 'with', 'this', 'that', 'from',
                    'into', 'over', 'when', 'what', 'where', 'each',
                    'just', 'like', 'some', 'time', 'come', 'came',
                    'gave', 'give', 'took', 'take', 'tell', 'told',
                    'know', 'knew', 'left', 'land', 'name', 'last',
                    'find', 'found', 'four', 'five', 'send', 'sent',
                    'next', 'much', 'more', 'most', 'good', 'long',
                    'years', 'year', 'days', 'day', 'son', 'sons',
                    'wife', 'wives', 'house', 'land'}

    if is_single_word and 2 <= len(text_lower) <= 25:
        if any(vp in text_lower for vp in vowel_patterns):
            return True
        for suf in suffixes:
            if text_lower.endswith(suf) and len(text_lower) > len(suf) + 1:
                return True
        if text_lower in common_short:
            return True
    elif not is_single_word:
        # Multi-word: count words that match English patterns
        words = text_lower.split()
        matches = 0
        for w in words:
            wclean = w.strip('.,;:!?\'"()')
            if not wclean: continue
            if wclean in common_short:
                matches += 1
                continue
            if 3 <= len(wclean) <= 25:
                if any(vp in wclean for vp in vowel_patterns):
                    matches += 1
                    continue
                if any(wclean.endswith(s) and len(wclean) > len(s) + 1
                       for s in suffixes):
                    matches += 1
                    continue
        # If majority of words look English, it's English
        if matches >= max(2, len(words) * 0.5):
            return True

    # English proper noun (transliterated names): ASCII-only with first letter
    # uppercase and 2+ lowercase letters following.  Davka Hebrew encoding 
    # rarely produces this pattern (it uses uppercase only for letters with dagesh,
    # which are followed by nikud bytes, not lowercase letters).
    if (3 <= len(text) <= 30 and text[0].isupper() and text[0].isalpha()):
        lower_count = sum(1 for c in text if c.islower())
        if lower_count >= 2:
            return True

    return False

def has_page_break(ev):
    return (ev['type'] == 'run'
            and ev['style'] in MISHNA_STYS
            and ev['bytes']
            and ev['bytes'][0] == 0x0c)

def is_heb(sty): return sty in HEB_STYLES


_HEB_SECTION_HEADING_RE = re.compile(
    r'\([א-ת"\s\-]+\)|\)[א-ת"\s\-]+\('  # parenthesized pasuk ref like "(א-ט)"
)


def looks_like_section_heading(text):
    """Detect if standalone text is a Hebrew section/chapter heading.

    A heading is plain Hebrew (no nikud) with a parenthesized verse
    reference, or contains 'פרק'/'פרשת'.  Typical examples:
      - "הליכת אברם לארץ כנען )א-ט("
      - "פרק יב"
      - "פרשת לכ לכ פרק י\"ב"
    """
    if not isinstance(text, str):
        return False
    text = text.strip()
    if len(text) < 5 or len(text) > 100:
        return False
    # No nikud or trup → it's not pesukim body text
    if any(0x0591 <= ord(c) <= 0x05c7 for c in text):
        return False
    # Has Hebrew letters
    if not any(0x05d0 <= ord(c) <= 0x05ea for c in text):
        return False
    # Pattern 1: parenthesized pasuk reference
    if _HEB_SECTION_HEADING_RE.search(text):
        return True
    # Pattern 2: explicit chapter or parsha keyword
    if 'פרק ' in text or 'פרשת ' in text:
        return True
    return False


def looks_like_section_heading_block(blk):
    """Detect if a TextBlock is predominantly a Hebrew section heading.

    A heading-block is short (< 60 chars), has 1-2 runs, no nikud,
    and is at least 80% Hebrew letters with a chapter-like reference.
    """
    if not hasattr(blk, 'runs') or not blk.runs:
        return False
    if len(blk.runs) > 2:
        return False
    full_text = ''.join(t for _, t in blk.runs)
    if len(full_text) < 5 or len(full_text) > 60:
        return False
    # No nikud/trup
    if any(0x0591 <= ord(c) <= 0x05c7 for c in full_text):
        return False
    # Predominantly Hebrew
    heb_letters = sum(1 for c in full_text if 0x05d0 <= ord(c) <= 0x05ea)
    eng_letters = sum(1 for c in full_text if c.isalpha() and c.isascii())
    total = heb_letters + eng_letters
    if total == 0 or heb_letters < total * 0.8:
        return False
    # Has parenthesized pasuk reference or chapter keyword
    if _HEB_SECTION_HEADING_RE.search(full_text):
        return True
    if 'פרק ' in full_text or 'פרשת ' in full_text:
        return True
    return False


_CHAPTER_START_RE = re.compile(r'^פרק\s+[אבגדהוזחטיכלמנסעפצקרשת]')


def is_chapter_start_block(blk):
    """Detect if a body block starts with 'פרק <letter>' (chapter beginning).

    Used to insert page breaks before each Bible chapter section.
    """
    if not hasattr(blk, 'runs') or not blk.runs:
        return False
    first_text = blk.runs[0][1].strip()
    return bool(_CHAPTER_START_RE.match(first_text))


def is_subtitle_pair_block(blk):
    """Detect if block is an English summary + Hebrew chapter heading pair.

    Pattern: 2 runs, one English-only and one Hebrew-only (no nikud), where
    the Hebrew run looks like a section heading.  These should render as
    a centered subtitle pair (English line, then Hebrew line below).
    """
    if not hasattr(blk, 'runs') or not blk.runs:
        return False
    if len(blk.runs) != 2:
        return False
    txt_a = blk.runs[0][1]
    txt_b = blk.runs[1][1]
    # No nikud anywhere
    if any(0x0591 <= ord(c) <= 0x05c7 for c in txt_a + txt_b):
        return False
    a_heb = any(0x0590 <= ord(c) <= 0x05ff for c in txt_a)
    b_heb = any(0x0590 <= ord(c) <= 0x05ff for c in txt_b)
    a_eng = any(c.isalpha() and c.isascii() for c in txt_a)
    b_eng = any(c.isalpha() and c.isascii() for c in txt_b)
    # One must be English-only and one Hebrew-only
    if not ((a_eng and not a_heb and b_heb and not b_eng)
            or (a_heb and not a_eng and b_eng and not b_heb)):
        return False
    # Reasonable total length
    full = txt_a + txt_b
    if len(full) > 120 or len(full) < 10:
        return False
    # Hebrew portion has a heading-like reference
    heb_text = txt_a if a_heb else txt_b
    if _HEB_SECTION_HEADING_RE.search(heb_text):
        return True
    if 'פרק ' in heb_text or 'פרשת ' in heb_text:
        return True
    return False



# ── Document model ────────────────────────────────────────────────────────────
class Block: pass

class TextBlock(Block):
    def __init__(self, role):
        self.role = role
        self.runs = []
    @property
    def text(self): return ''.join(t for _,t in self.runs)
    def add(self, s, t):
        if t: self.runs.append((s, t))

class KeyWordBlock(Block):
    def __init__(self):
        self.pairs = []

class ImageBlock(Block):
    """An embedded image extracted from the DWD file."""
    def __init__(self, raw, width, height, fmt='JPEG', index=0):
        self.raw    = raw
        self.width  = width
        self.height = height
        self.fmt    = fmt
        self.index  = index

class ParshaTopicsBlock(Block):
    """Parsha summary chart: rows of (num, topic_heb, topic_eng, points, pesukim)."""
    def __init__(self):
        self.rows = []   # list of dicts

class QABlock(Block):
    """Numbered Q&A list: list of (question_runs, answer_runs) pairs."""
    def __init__(self, inyan_title=''):
        self.inyan_title = inyan_title
        self.items = []  # list of {'q': [(sty,text)...], 'a': [(sty,text)...]}

class PronounBlock(Block):
    """Pronoun/verb conjugation chart."""
    def __init__(self):
        # Top 7-column grid headers
        self.heb_pronouns  = []   # Hebrew pronouns (אני אתה הוא ...)
        self.eng_pronouns  = []   # English labels (I You He ...)
        self.prefix_forms  = []   # לשרש forms per pronoun
        self.shoresh_forms = []   # שרש forms per pronoun
        self.obj_suffixes  = []   # object suffixes (אתי אתך ...)
        self.obj_labels    = []   # English object labels (Me You Him ...)
        # Bottom conjugation lists  [(label, [(heb, eng), ...])]
        self.sections = []   # list of (box_label, [(heb_form, eng_label), ...])

class SongStanzaBlock(Block):
    """A song stanza with a style label (Low/High) and lines."""
    def __init__(self, style_label='', title=''):
        self.style_label = style_label   # 'Low', 'High', etc.
        self.title = title
        self.lines = []  # list of (sty, text)


def _classify_table_region(events, start, end, decode):
    """Examine a window of events and decide if they form a known table type.

    Returns (table_type, extra_info) or (None, None).
    Table types: 'parsha_topics', 'qa', 'pronoun'.
    """
    styles = {ev['style'] for ev in events[start:end] if ev['type'] == 'run'}
    texts  = [decode(ev) for ev in events[start:end] if ev['type'] == 'run']

    # ── Parsha Topics signature ───────────────────────────────────────────────
    # Has a numeric row-number column and a "topic" Hebrew column + points column
    # Detect: row-number style (single digit), topic style, description styles
    # Heuristic: has both a style that only ever produces short digits AND
    # a style that produces Hebrew topic names (1-3 words)
    digit_runs = [t for t in texts if t.strip().isdigit()]
    if len(digit_runs) >= 3:
        # Could be parsha topics - check for multi-run mixed Hebrew/English lines
        mixed = [t for t in texts if t.strip() and not t.strip().isdigit()
                 and len(t.strip()) > 5]
        if len(mixed) >= 3:
            return 'parsha_topics', {}

    # ── Q&A signature ────────────────────────────────────────────────────────
    # Long English questions followed by answers; triggered by '?' in question text
    questions = [t for t in texts if '?' in t and len(t.strip()) > 10]
    if len(questions) >= 2:
        return 'qa', {}

    # ── Pronoun chart signature ───────────────────────────────────────────────
    # Has "אני", "אתה", "הוא" and "I", "You", "He" etc. in close proximity
    PRONOUNS_HEB = {'אני', 'אתה', 'הוא', 'היא', 'אנחנו', 'אתם', 'הם',
                    'אֲנִי', 'אַתָה', 'הוא', 'הִיא'}
    PRONOUNS_ENG = {'I', 'You', 'He', 'She', 'We', 'They', 'Me', 'Him', 'Her'}
    heb_matches = sum(1 for t in texts if t.strip() in PRONOUNS_HEB)
    eng_matches = sum(1 for t in texts if t.strip() in PRONOUNS_ENG)
    if heb_matches >= 3 and eng_matches >= 3:
        return 'pronoun', {}

    return None, None


def build_model(events, with_nikud=True, with_trup=True):
    blocks = []
    cur_text    = None
    cur_kw      = None
    in_kw_sec   = False
    kw_line_heb = []
    kw_line_eng = []
    para_gap    = 0
    heading_count = 0

    def dec(ev):
        return decode_run(ev, with_nikud, with_trup)

    def flush_text():
        nonlocal cur_text
        if cur_text and cur_text.runs:
            blocks.append(cur_text)
        cur_text = None

    def flush_kw_line():
        nonlocal kw_line_heb, kw_line_eng
        heb = ' '.join(h for h in kw_line_heb if h)
        eng = ' '.join(e for e in kw_line_eng if e)
        if heb or eng:
            cur_kw.pairs.append((heb, eng))
        kw_line_heb = []; kw_line_eng = []

    def flush_kw():
        nonlocal cur_kw, in_kw_sec, kw_line_heb, kw_line_eng
        if in_kw_sec: flush_kw_line()
        if cur_kw and cur_kw.pairs: blocks.append(cur_kw)
        cur_kw = None; in_kw_sec = False; kw_line_heb = []; kw_line_eng = []

    img_index = [0]

    # ── Build a per-event text cache for lookahead ─────────────────────────
    _text_cache = {}
    def get_text(ev):
        eid = id(ev)
        if eid not in _text_cache:
            if ev['type'] != 'run':
                _text_cache[eid] = ''
            else:
                t = dec(ev)
                _text_cache[eid] = '' if _is_marker_run(t) else t
        return _text_cache[eid]

    # ── Table/section detection helpers ───────────────────────────────────
    def styles_in_window(start, length=60):
        """Collect styles from a window of events ahead."""
        s = set()
        for ev in events[start:start+length]:
            if ev['type'] == 'run':
                s.add(ev['style'])
        return s

    def texts_in_window(start, length=60):
        return [get_text(ev) for ev in events[start:start+length]
                if ev['type'] == 'run' and get_text(ev).strip()]

    # ── Parse a Parsha Topics block (4-col table) starting at index i ──────
    def parse_parsha_topics(i):
        """Parse the parsha summary chart from event index i.
        Returns (ParshaTopicsBlock, new_i).
        Layout per row: row_num | topic_heb | topic_eng | (description lines) | points | pesukim
        In the event stream, each row looks like:
          [row_num(digit)] PARA [topic_heb] PARA [description runs...] PARA [points] PARA [pesukim] PARA [next row_num...]
        The row_num style produces single digits. Topic_heb style produces short Hebrew.
        """
        blk = ParshaTopicsBlock()
        # Detect the row-number style: find first run that's a single-digit '1'
        # (start of first row). Look ahead up to 25 events past column headers.
        row_num_sty = None
        for ev in events[i:i+25]:
            if ev['type'] == 'run':
                t = get_text(ev)
                if t.strip() == '1':   # first row must be '1'
                    row_num_sty = ev['style']
                    break
        if row_num_sty is None:
            return None, i

        cur_row = None
        j = i
        # Termination: stop when we see another section header (style 0x07/0x30/0xc2/0xbe)
        # or more than 6 blank PARAₛ in a row (section boundary)
        # Actually use a simpler stop: stop when row_num exceeds 9 (more than 9 inyanim unlikely)
        consec_paras = 0
        while j < len(events):
            ev = events[j]
            if ev['type'] == 'image':
                break
            if ev['type'] == 'para':
                consec_paras += 1
                # After 8+ consecutive paras without content, we've left the table
                if consec_paras > 8 and cur_row and cur_row['desc']:
                    blk.rows.append(cur_row)
                    cur_row = None
                    break
                j += 1; continue
            if ev['type'] != 'run':
                j += 1; continue

            t = get_text(ev)
            if not t.strip():
                j += 1; continue
            if _is_marker_run(t):
                j += 1; continue
            consec_paras = 0

            sty = ev['style']

            # New row starts with a digit in the row_num_sty
            if sty == row_num_sty and t.strip().isdigit():
                if cur_row:
                    blk.rows.append(cur_row)
                cur_row = {'num': t.strip(), 'topic_heb': '', 'topic_eng': '',
                           'desc': [], 'points': '', 'pesukim': ''}
                j += 1; continue

            if cur_row is None:
                # Skip header rows (non-digit, before first row)
                j += 1; continue

            # After we have a row, classify subsequent runs
            # Topic Hebrew: first Hebrew run after digit (style 0x65)
            if not cur_row['topic_heb'] and ev.get('is_hebrew', _is_hebrew_content(ev['bytes'])) and not cur_row['desc']:
                cur_row['topic_heb'] = t.strip()
                j += 1; continue

            # Points: standalone digit at the end of a row (comes after description)
            if t.strip().isdigit() and sty != row_num_sty and cur_row['desc']:
                if not cur_row['points']:
                    cur_row['points'] = t.strip()
                j += 1; continue

            # Pesukim reference: contains פרק or " (comes after points)
            if cur_row.get('points') and not cur_row['pesukim']:
                if 'פרק' in t or ('"' in t and len(t.strip()) < 20):
                    cur_row['pesukim'] = t.strip()
                    j += 1; continue

            # Everything else is description content
            cur_row['desc'].append((sty, t))
            j += 1

        if cur_row and (cur_row['desc'] or cur_row['topic_heb']):
            blk.rows.append(cur_row)
        return blk, j

    # ── Parse Q&A section starting at index i ──────────────────────────────
    def parse_qa(i, inyan_title=''):
        """Parse numbered Q&A section using style-based separation.

        From inspection of the Lech Lecha file:
          - Question runs use a small "question style" (e.g. 0x36 Arial 8pt)
          - Inline Hebrew words within questions use q_style - 2 (e.g. 0x34)
          - Answer runs use other, larger styles (0x35, 0x37, 0x3b, 0x32, 0x33)
          - Each Q&A item ends at a PARA break

        Strategy:
          1. Detect q_style from the first run that ends with '?'
          2. Collect runs: q_style and q_style-2 go to question; others to answer
          3. A PARA break flushes the item
        """
        blk = QABlock(inyan_title)

        # Detect question style: first run ending with '?' or '?\r'
        q_style = None
        for ev in events[i:i+100]:
            if ev['type'] == 'run':
                t = get_text(ev).rstrip('\r\n ')
                if t.endswith('?'):
                    q_style = ev['style']
                    break
        if q_style is None:
            return None, i

        # Inline-Hebrew-in-question style (typically q_style - 2)
        q_inline_style = q_style - 2

        j = i
        cur_q = []
        cur_a = []
        seen_q_mark = False     # have we seen '?' in the current question?

        # Stop markers
        end_markers = {'Who Did It?', 'כינוי לגוף הפועל', 'כותרות הפרשה'}

        def flush_item():
            nonlocal cur_q, cur_a, seen_q_mark
            if cur_q or cur_a:
                blk.items.append({'q': cur_q, 'a': cur_a})
            cur_q = []; cur_a = []; seen_q_mark = False

        while j < len(events):
            ev = events[j]

            if ev['type'] == 'image':
                flush_item()
                break

            if ev['type'] == 'para':
                flush_item()
                j += 1; continue

            if ev['type'] != 'run':
                j += 1; continue

            t = get_text(ev)
            if not t.strip():
                j += 1; continue
            if _is_marker_run(t):
                j += 1; continue

            sty = ev['style']

            # End the section when we hit another major header
            if t.strip() in end_markers:
                flush_item()
                return blk, j

            # Classify this run
            if sty == q_style or (sty == q_inline_style and not seen_q_mark):
                cur_q.append((sty, t))
                if '?' in t:
                    seen_q_mark = True
            else:
                # Non-question style: goes to answer
                cur_a.append((sty, t))

            j += 1

        flush_item()
        return blk, j

    # ── Parse Pronoun chart starting at index i ─────────────────────────────
    def parse_pronoun(i):
        """Parse the pronoun/verb conjugation chart.

        Structure (verified from events + PDF):

        SECTION 1 - 'Who Did It?' (Subject):
          Top 7-column header grid:
            Row 1: Hebrew pronouns (style 0x7c, space-separated)
                   אֲנִי  אַתָה  הוא  הִיא  אֲנַחְנו  אַתֶם  הֵם
            Row 2: English labels (style 0x0f, space-separated)
                   I  You  He  She  We  You All  They
                   (note: text appears in RTL order when displayed)
            Rows 3+: verb prefix letters, suffix letters (style 0x7c, individual tokens)

          Paired conjugation lists (style 0x0b = Heb, 0x31 = Eng):
            Multiple 7-item groups, each group = one conjugation paradigm
            Separated by double-PARA breaks

        SECTION 2 - 'To Whom Was It Done' (Object):
          Header: 'Mipei lbes dRrEl' (style 0xd5)
          Similar structure with object suffixes

        Termination: stop at next major section header (0xd3 with new title,
        image, or unrelated content).
        """
        blk = PronounBlock()
        j = i

        # Styles in the pronoun chart family
        PRONOUN_STYLES = {0x7c, 0x0f, 0x78, 0x0b, 0x31, 0xd3, 0xd5}

        # State tracking
        in_pairs = False          # Are we in the 0x0b/0x31 pair lists?
        cur_list_pairs = []       # Current paradigm being collected
        cur_list_label = ''       # Label/header for current list
        consecutive_paras = 0

        HPRON_BASE = {'אני','אתה','הוא','היא','אנחנו','אתם','הם'}
        EPRON_BASE = {'I','You','He','She','We','They','You all','You All','Y ou'}

        def _strip_nikud(s):
            return ''.join(c for c in s if ord(c) < 0x0591 or ord(c) > 0x05C7)

        def _is_pronoun_heb(t):
            bare = _strip_nikud(t.strip())
            return bare in HPRON_BASE

        def _is_pronoun_eng(t):
            return t.strip() in EPRON_BASE

        def _finish_list():
            nonlocal cur_list_pairs, cur_list_label
            if cur_list_pairs:
                blk.sections.append((cur_list_label, cur_list_pairs))
                cur_list_pairs = []
                cur_list_label = ''

        while j < len(events):
            ev = events[j]

            if ev['type'] == 'image':
                break

            if ev['type'] == 'para':
                consecutive_paras += 1
                # Double-PARA breaks a paradigm list
                if consecutive_paras >= 2 and cur_list_pairs:
                    _finish_list()
                j += 1; continue

            if ev['type'] != 'run':
                j += 1; continue

            t = get_text(ev)
            if not t.strip():
                j += 1; continue
            if _is_marker_run(t):
                j += 1; continue

            sty = ev['style']

            # Stop if we've moved out of the pronoun chart area
            if sty not in PRONOUN_STYLES:
                # ... unless we're clearly still in a table (e.g. Bible reference)
                if sty != 0x0b and sty != 0x31:
                    break

            consecutive_paras = 0
            is_h = ev.get('is_hebrew', _is_hebrew_content(ev['bytes']))

            # Section header — 'Who Did It?' or 'To Whom Was It Done'
            if sty == 0xd3:
                _finish_list()
                cur_list_label = t.strip()
                j += 1; continue

            if sty == 0xd5:
                _finish_list()
                cur_list_label = t.strip()
                j += 1; continue

            # Top-grid Hebrew pronouns (space-separated in one run)
            if sty == 0x7c and _is_pronoun_heb(t.split()[0] if t.split() else ''):
                for tok in t.split():
                    if tok.strip() and _is_pronoun_heb(tok):
                        blk.heb_pronouns.append((sty, tok.strip()))
                j += 1; continue

            # Top-grid English labels (space-separated)
            if sty == 0x0f:
                # Could be one token or many
                # If comma/space separated with pronoun tokens, they're labels
                toks = [tok.strip() for tok in t.split() if tok.strip()]
                # Detect if this is pronoun row vs object row based on content
                has_obj = any(tok in {'Me','Him','Her','Us','Them'} for tok in toks)
                for tok in toks:
                    if tok in {'Me','Him','Her','Us','Them'} or (has_obj and tok == 'You'):
                        blk.obj_labels.append((sty, tok))
                    elif tok in EPRON_BASE:
                        if len(blk.eng_pronouns) < 7:
                            blk.eng_pronouns.append((sty, tok))
                j += 1; continue

            # Other 0x7c runs — verb prefix/suffix letter rows (not pronouns)
            # Skip these individual letter tokens; they're grid cells best shown as grid
            if sty == 0x7c:
                # These are best rendered as additional grid rows but we can skip for now
                j += 1; continue

            # Object suffix forms (style 0x78)
            if sty == 0x78:
                # These pair with obj_labels; keep them
                blk.sections.append(('object_suffixes', [(sty, t.strip(), True)]))
                j += 1; continue

            # Paired conjugation (0x0b Hebrew + 0x31 English)
            if sty == 0x0b:
                cur_list_pairs.append({'heb': t.strip(), 'eng': ''})
                j += 1; continue

            if sty == 0x31:
                if cur_list_pairs and not cur_list_pairs[-1]['eng']:
                    cur_list_pairs[-1]['eng'] = t.strip().rstrip('\r')
                j += 1; continue

            j += 1

        _finish_list()
        return blk, j

    # ── Main event loop (index-based for table lookahead) ─────────────────
    i = 0
    while i < len(events):
        ev = events[i]

        if ev['type'] == 'image':
            flush_text(); flush_kw()
            blocks.append(ImageBlock(ev['raw'], ev['width'], ev['height'],
                                     ev.get('fmt','JPEG'), img_index[0]))
            img_index[0] += 1
            i += 1; continue

        if ev['type'] == 'para':
            para_gap += 1
            if in_kw_sec:
                flush_kw_line()
            else:
                if para_gap >= 2:
                    flush_text()
                    if blocks and not (isinstance(blocks[-1], TextBlock)
                                       and blocks[-1].role == 'blank'):
                        blocks.append(TextBlock('blank'))
                    para_gap = 0
                elif cur_text and cur_text.runs:
                    flush_text()
            i += 1; continue

        if ev['type'] != 'run':
            i += 1; continue

        para_gap = 0
        sty  = ev['style']
        text = get_text(ev)
        if sty in (SECTION_HDR_STY, HEB_HEADING_STY):
            text = text.strip()
        if not text:
            i += 1; continue

        # Markers already filtered in get_text(), but double-check
        if _is_marker_run(text):
            i += 1; continue

        # ── Table / section detection ─────────────────────────────────────
        # 'כותרות הפרשה' = inyan title list (English+Hebrew heading pairs)
        # Render as a formatted header + body text
        if text.strip() == 'כותרות הפרשה':
            flush_text(); flush_kw()
            # Emit the section header
            hdr_blk = TextBlock('section_hdr'); hdr_blk.add(sty, 'Parsha Topics  /  כותרות הפרשה')
            blocks.append(hdr_blk)
            # Consume inyan title runs until we hit '621 Rqewim' or 'Topic'
            j = i + 1
            cur_body = None
            cur_eng = ''; cur_heb = ''
            while j < len(events):
                ev2 = events[j]
                if ev2['type'] == 'para': j += 1; continue
                if ev2['type'] != 'run': j += 1; continue
                t2 = get_text(ev2)
                if not t2.strip() or _is_marker_run(t2): j += 1; continue
                if '621' in t2 or t2.strip() == 'Topic': break
                # Emit each run as a body block
                if cur_body is None:
                    cur_body = TextBlock('body')
                is_h2 = ev2.get('is_hebrew', _is_hebrew_content(ev2['bytes']))
                cur_body.add(ev2['style'], t2)
                j += 1
                # At Hebrew run, flush as a line
                if is_h2 and cur_body.runs:
                    blocks.append(cur_body)
                    cur_body = None
            if cur_body and cur_body.runs:
                blocks.append(cur_body)
            i = j  # position now at '621 Rqewim' or 'Topic'
            continue

        # '621 Rqewim' = start of the 7-row parsha summary chart
        if '621' in text or '126' in text:
            flush_text(); flush_kw()
            blk, i = parse_parsha_topics(i + 1)
            if blk and blk.rows:
                blocks.append(blk)
            continue

        # Q&A section
        if 'שאלות' in text and 'תשובות' in text:
            flush_text(); flush_kw()
            blk, i = parse_qa(i + 1, inyan_title=text.strip())
            if blk and blk.items:
                blocks.append(blk)
            continue

        # Pronoun chart (header line triggers it)
        if 'Who Did It?' in text or 'כינוי לגוף הפועל' in text:
            flush_text(); flush_kw()
            blk, i = parse_pronoun(i)
            if blk and (blk.heb_pronouns or blk.sections):
                blocks.append(blk)
            continue

        if has_page_break(ev):
            flush_text(); flush_kw()
            blocks.append(TextBlock('page_break'))
            rest = decode_run({'type':'run','style':sty,
                               'bytes': ev['bytes'][1:],
                               'use_style': ev.get('use_style', False)},
                              with_nikud, with_trup).strip()
            if rest:
                cur_text = TextBlock('mishna'); cur_text.add(sty, rest)
            i += 1; continue

        if sty == SECTION_HDR_STY:
            flush_text(); flush_kw()
            hdr = TextBlock('section_hdr'); hdr.add(sty, text)
            blocks.append(hdr)
            if text == 'KEY WORDS':
                in_kw_sec = True; cur_kw = KeyWordBlock()
            i += 1; continue

        if in_kw_sec:
            if sty == HEB_HEADING_STY:
                flush_kw()
            else:
                if is_heb(sty) or _is_hebrew_content(ev['bytes']):
                    kw_line_heb.append(text)
                else:
                    kw_line_eng.append(text)
                i += 1; continue

        if sty == HEB_HEADING_STY:
            flush_text(); flush_kw()
            # Only auto-insert page break for Format A (use_style format), where
            # HEB_HEADING_STY genuinely marks chapter starts.  In other formats
            # the same style byte is used for short song fragments which
            # shouldn't trigger page breaks.
            if heading_count > 0 and ev.get('use_style', False):
                blocks.append(TextBlock('page_break'))
            heading_count += 1
            tb = TextBlock('heading'); tb.add(sty, text)
            blocks.append(tb)
            i += 1; continue

        if sty in MISHNA_STYS:
            if not cur_text or cur_text.role != 'mishna':
                flush_text(); cur_text = TextBlock('mishna')
            cur_text.add(sty, text)
            i += 1; continue

        if not cur_text or cur_text.role != 'body':
            flush_text(); cur_text = TextBlock('body')
        cur_text.add(sty, text)
        i += 1

    flush_text(); flush_kw()
    return blocks


# ── DOCX XML helpers ──────────────────────────────────────────────────────────
def _pPr(p): return p._p.get_or_add_pPr()

def _bidi(p, on=True):
    pr = _pPr(p)
    el = OxmlElement('w:bidi'); el.set(qn('w:val'), '1' if on else '0'); pr.append(el)
    if on:
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'right'); pr.append(jc)

def _spacing(p, before=0, after=0, line=None):
    pr = _pPr(p)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before)); sp.set(qn('w:after'), str(after))
    if line: sp.set(qn('w:line'), str(line)); sp.set(qn('w:lineRule'), 'auto')
    pr.append(sp)

def _border_bottom(p, color='1F4E79', sz=12):
    pr = _pPr(p); pb = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
    b.set(qn('w:space'),'2');    b.set(qn('w:color'),color)
    pb.append(b); pr.append(pb)

def _box_border(p, color='2E75B6', fill='DBE9F9'):
    pr = _pPr(p); pb = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
        b.set(qn('w:space'),'4');    b.set(qn('w:color'),color)
        pb.append(b)
    pr.append(pb)
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),fill)
    pr.append(sh)

def _rtl_run(para, text, font='David', sz=13, bold=False, color=None):
    r = para.add_run(text); r.font.name = font; r.font.size = Pt(sz)
    if bold:  r.font.bold = True
    if color: r.font.color.rgb = color
    rp = r._r.get_or_add_rPr()
    rt = OxmlElement('w:rtl'); rt.set(qn('w:val'),'1'); rp.append(rt)
    rf = rp.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rp.insert(0, rf)
    rf.set(qn('w:cs'), font)
    return r

def _ltr_run(para, text, font='Calibri', sz=11, bold=False, color=None):
    r = para.add_run(text); r.font.name = font; r.font.size = Pt(sz)
    if bold:  r.font.bold = True
    if color: r.font.color.rgb = color
    return r

def _inline_heb(para, text, sz=12, bold=False):
    r = para.add_run(text); r.font.name = 'David'; r.font.size = Pt(sz)
    if bold: r.font.bold = True
    rp = r._r.get_or_add_rPr()
    rt = OxmlElement('w:rtl'); rt.set(qn('w:val'),'1'); rp.append(rt)
    rf = rp.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rp.insert(0, rf)
    rf.set(qn('w:cs'), 'David')
    return r


# ── KEY WORDS list (inline format) ────────────────────────────────────────────
def _add_kw_table(doc, pairs):
    """Render Key Words in original PDF inline format:
       הַסַּפָּר )א( - the barber
    Each entry is one paragraph with Hebrew letter in parens, then dash, then English.
    """
    _HEB_LETTERS_KW = ['א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט', 'י',
                       'יא', 'יב', 'יג', 'יד', 'טו', 'טז', 'יז', 'יח', 'יט', 'כ']

    for idx, (heb, eng) in enumerate(pairs):
        p = doc.add_paragraph()
        _bidi(p, False)
        _spacing(p, before=30, after=20)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.3)

        letter = _HEB_LETTERS_KW[idx] if idx < len(_HEB_LETTERS_KW) else str(idx + 1)

        # Hebrew word (RTL embedded)
        if heb.strip():
            heb_run = _rtl_run(p, heb.strip(), font='David', sz=13, bold=True)
        # Hebrew letter in parens like ")א("  (Note: parens around Hebrew letter
        # appear reversed visually in BiDi rendering)
        paren_run = p.add_run(f' )‏{letter}‏( ')
        paren_run.font.name = 'David'
        paren_run.font.size = Pt(12)
        # Dash and English translation
        eng_text = eng.strip().lstrip('-').strip()
        if eng_text:
            dash_run = p.add_run('- ')
            dash_run.font.name = 'Times New Roman'
            dash_run.font.size = Pt(12)
            eng_run = p.add_run(eng_text)
            eng_run.font.name = 'Times New Roman'
            eng_run.font.size = Pt(12)

    doc.add_paragraph()


# ── DOCX builder ─────────────────────────────────────────────────────────────
FONT_NOTE = (
    "Font note: This document uses the David font (built into Windows) as fallback. "
    "For authentic Davka Stam appearance, install one of these free fonts:\n"
    "  • Stam Ashkenaz CLM  →  hebrewfont.net  (best Davka Stam replacement; full nikud/trup)\n"
    "  • Frank Ruehl CLM    →  hebrewfont.net  (classic Israeli serif body font)\n"
    "  • Ezra SIL            →  software.sil.org/ezrasil/  (scholarly; full cantillation)\n"
    "  • Noto Serif Hebrew   →  fonts.google.com  (modern open serif)\n"
    "After installing: Ctrl+A to select all, then choose the font in the Home tab."
)
NAVY  = RGBColor(0x1F, 0x4E, 0x79)
GRAY  = RGBColor(0x50, 0x50, 0x50)

# Fallback sizes when the style table has size_pt == 0
_FALLBACK_HEB_PT  = 13
_FALLBACK_ENG_PT  = 11

_HEBREW_FONTS = frozenset([
    'David', 'Frank Ruehl CLM', 'Noto Serif Hebrew', 'Noto Sans Hebrew',
    'Noto Rashi Hebrew', 'Ezra SIL', 'SBL Hebrew',
    # Also accept Davka- prefixes if they slipped through
    'Davka Stam', 'Davka FrankRuhl', 'Davka Hadasah', 'Davka David',
    'Davka Rashi', 'Davka Kastel', 'Davka Yerushalmy', 'Davka Meir',
    'Davka Siddur', 'Davka Drogolin',
])
_ENGLISH_FONTS = frozenset([
    'Arial', 'Times New Roman', 'Calibri', 'Courier New',
    'Cambria', 'Georgia', 'Verdana',
])

def _validate_font_for_content(font_name, size_pt, is_hebrew):
    """If the style table's font doesn't match the run's actual content
    language, override to a sensible default. This handles DWD files where
    the CDocStyle indexing doesn't match run style bytes exactly."""
    font_is_heb = font_name in _HEBREW_FONTS
    font_is_eng = font_name in _ENGLISH_FONTS
    if is_hebrew and font_is_eng:
        # Content is Hebrew but style says English font → use David at same size
        return 'David', size_pt
    if (not is_hebrew) and font_is_heb:
        # Content is English but style says Hebrew font → use Arial at reasonable size
        # If the size came from a Hebrew style, it may be way too large for English
        # Cap at 14pt for safety since English typically shouldn't be larger
        eng_size = min(size_pt, 14.0) if size_pt > 14 else size_pt
        return 'Arial', eng_size
    return font_name, size_pt


def _styled_run(para, text, props, is_hebrew):
    """Add a run to para, applying formatting from a parsed style props dict.

    If the style's font conflicts with the content's language (e.g. content
    is clean English but the style uses a Hebrew font), override to a sensible
    default.  This makes the converter robust to CDocStyle indexing mismatches.
    """
    font_name = props.get('font', 'David' if is_hebrew else 'Calibri')
    size_pt   = props.get('size_pt', 0) or (_FALLBACK_HEB_PT if is_hebrew else _FALLBACK_ENG_PT)
    bold      = props.get('bold', False)
    italic    = props.get('italic', False)
    underline = props.get('underline', False)

    # Validate that font matches content type
    font_name, size_pt = _validate_font_for_content(font_name, size_pt, is_hebrew)

    r = para.add_run(text)
    r.font.name      = font_name
    r.font.size      = Pt(size_pt)
    r.font.bold      = bold      or None
    r.font.italic    = italic    or None
    r.font.underline = underline or None

    if is_hebrew:
        rp = r._r.get_or_add_rPr()
        rt = OxmlElement('w:rtl'); rt.set(qn('w:val'), '1'); rp.append(rt)
        rf = rp.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rp.insert(0, rf)
        rf.set(qn('w:cs'), font_name)
    return r


def _render_parsha_topics(doc, blk, style_table):
    """Render the parsha summary chart (7-row table) as a Word table.
    4 columns: # | Hebrew Topic | Description (mixed) | Points + Ref
    """
    if not blk.rows:
        return
    NAVY = RGBColor(0x1F, 0x4E, 0x79)

    # Title row
    title_p = doc.add_paragraph()
    _bidi(title_p, False); _spacing(title_p, before=160, after=60)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ltr_run(title_p, '126 ', sz=16, bold=True, color=NAVY)
    _rtl_run(title_p, 'פְּסוּקִים', font='David', sz=16, bold=True, color=NAVY)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl_xml.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Column widths in twips: # | Topic | Description | Ref
    col_widths = ['540', '1440', '5940', '1080']

    def _set_cell_width(cell, w):
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None: tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), w); tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    # Header row
    hdr = tbl.rows[0]
    headers = [('#', False), ('TOPIC / כותרות', True), ('POINTS', False), ('# פסוקים', True)]
    for cell, (txt, is_h), w in zip(hdr.cells, headers, col_widths):
        _set_cell_width(cell, w)
        p = cell.paragraphs[0]
        _bidi(p, is_h); _spacing(p, before=40, after=40)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_h:
            _rtl_run(p, txt, font='David', sz=10, bold=True, color=NAVY)
        else:
            _ltr_run(p, txt, sz=10, bold=True, color=NAVY)

    # Data rows
    for row_data in blk.rows:
        row = tbl.add_row()
        cells = row.cells

        for cell, w in zip(cells, col_widths):
            _set_cell_width(cell, w)

        # Col 0: row number
        p0 = cells[0].paragraphs[0]
        _bidi(p0, False); _spacing(p0, before=60, after=60)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _ltr_run(p0, row_data.get('num', ''), sz=16, bold=True)

        # Col 1: Hebrew topic
        p1 = cells[1].paragraphs[0]
        _bidi(p1, True); _spacing(p1, before=40, after=40)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        topic = row_data.get('topic_heb', '') or row_data.get('topic_eng', '')
        if topic:
            is_h = _is_hebrew_content(topic.encode('utf-8', 'ignore'))
            if is_h:
                _rtl_run(p1, topic, font='David', sz=14, bold=True)
            else:
                _ltr_run(p1, topic, sz=11, bold=True)

        # Col 2: Description (mixed Hebrew/English runs)
        p2 = cells[2].paragraphs[0]
        _bidi(p2, False); _spacing(p2, before=40, after=40)
        for sty, txt in row_data.get('desc', []):
            props = dict(style_table.get(sty, {}))
            # Cap size for description column - tables need smaller text
            if props.get('size_pt', 0) > 11:
                props['size_pt'] = 11
            # Don't bold descriptions - they're body text
            props['bold'] = False
            run_is_heb = props.get('is_hebrew', _is_hebrew_content(txt.encode('utf-8', 'ignore')))
            _styled_run(p2, txt, props, run_is_heb)

        # Col 3: Points count + pesukim ref
        p3 = cells[3].paragraphs[0]
        _bidi(p3, False); _spacing(p3, before=40, after=40)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pts = row_data.get('points', '')
        ref = row_data.get('pesukim', '')
        if pts:
            r = p3.add_run(pts); r.font.size = Pt(13); r.font.bold = True
        if ref:
            r2 = p3.add_run('\n' + ref if pts else ref)
            r2.font.size = Pt(9)
            r2.font.name = 'David'
            rp = r2._r.get_or_add_rPr()
            rt = OxmlElement('w:rtl'); rt.set(qn('w:val'), '1'); rp.append(rt)

    doc.add_paragraph()


def _render_qa(doc, blk, style_table):
    """Render Q&A as a numbered list: question then bold answer."""
    NAVY = RGBColor(0x1F, 0x4E, 0x79)
    GRAY = RGBColor(0x50, 0x50, 0x50)

    def _classify(sty, txt):
        """Return True if run is Hebrew, applying content-trumps-style override."""
        text_bytes = txt.encode('utf-8', 'ignore')
        if any(0x0590 <= ord(c) <= 0x05ff for c in txt):
            return True
        if _is_clean_english(text_bytes):
            return False
        props = style_table.get(sty, {})
        return props.get('is_hebrew', False) or _is_hebrew_content(text_bytes)

    def _render_runs(p, runs, force_bold=False):
        """Render a sequence of (sty, text) runs into paragraph p,
        adding auto-space at Hebrew↔English boundaries."""
        prev_was_heb = None
        prev_ended_space = True
        for sty, txt in runs:
            run_is_heb = _classify(sty, txt)
            # Insert space at language boundary
            if (prev_was_heb is not None
                and prev_was_heb != run_is_heb
                and not prev_ended_space
                and txt and not txt[0].isspace()):
                p.add_run(' ')
            props = dict(style_table.get(sty, {}))
            if force_bold:
                props['bold'] = True
            _styled_run(p, txt, props, run_is_heb)
            prev_was_heb = run_is_heb
            prev_ended_space = bool(txt and txt[-1].isspace())

    # Inyan section header (e.g. "שאלות ותשובות פרשת לכ לכ")
    if blk.inyan_title:
        p = doc.add_paragraph()
        _spacing(p, before=180, after=60)
        _box_border(p)
        parts = blk.inyan_title.split()
        for word in parts:
            is_h = _is_hebrew_content(word.encode('utf-8', 'ignore'))
            if is_h:
                _bidi(p, True)
                _rtl_run(p, word + ' ', font='David', sz=12, bold=True, color=NAVY)
            else:
                _ltr_run(p, word + ' ', sz=11, bold=True, color=NAVY)

    for n, item in enumerate(blk.items, 1):
        # Question paragraph
        if item['q']:
            p = doc.add_paragraph()
            _bidi(p, False); _spacing(p, before=60, after=15)
            _ltr_run(p, f'.{n}  ', sz=11)
            _render_runs(p, item['q'], force_bold=False)

        # Answer paragraph (bold, indented)
        if item['a']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _spacing(p, before=5, after=55)

            heb_count = sum(1 for sty, txt in item['a']
                           if _is_hebrew_content(txt.encode('utf-8', 'ignore'))
                           or style_table.get(sty, {}).get('is_hebrew', False)
                           or any(0x0590 <= ord(c) <= 0x05ff for c in txt))
            para_is_heb = heb_count > len(item['a']) / 2
            _bidi(p, para_is_heb)
            if para_is_heb:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            _render_runs(p, item['a'], force_bold=True)

    doc.add_paragraph()


def _render_pronoun(doc, blk, style_table):
    """Render pronoun/verb conjugation chart.

    Structure:
      - 7-col header grid (Hebrew pronouns + English labels)
      - Multiple conjugation paradigms, each rendered as a labelled 2-col table
        (Hebrew form | English label), 7 rows per paradigm
    """
    NAVY = RGBColor(0x1F, 0x4E, 0x79)

    # Title
    title_p = doc.add_paragraph()
    _bidi(title_p, True); _spacing(title_p, before=120, after=40)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl_run(title_p, 'כִּינּוּי לְגוּף הַפּוֹעֵל', font='David', sz=14, bold=True, color=NAVY)
    _ltr_run(title_p, '   Who Did It? [Subject]', sz=11, color=NAVY)

    # Top 7-column subject pronoun header grid
    if blk.heb_pronouns:
        n = len(blk.heb_pronouns)
        tbl = doc.add_table(rows=2, cols=n)
        tbl.style = 'Table Grid'

        for col_idx, (sty, heb) in enumerate(blk.heb_pronouns):
            cell = tbl.rows[0].cells[col_idx]
            p = cell.paragraphs[0]
            _bidi(p, True); _spacing(p, before=40, after=20)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            props = style_table.get(sty, {'font': 'David', 'size_pt': 14})
            _styled_run(p, heb, props, True)

        # English labels appear in RTL order in source — reverse to align with Hebrew
        eng_ordered = list(reversed(blk.eng_pronouns[:n])) if blk.eng_pronouns else []
        for col_idx, (sty, eng) in enumerate(eng_ordered):
            cell = tbl.rows[1].cells[col_idx]
            p = cell.paragraphs[0]
            _bidi(p, False); _spacing(p, before=10, after=20)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            props = style_table.get(sty, {'font': 'Calibri', 'size_pt': 10})
            _styled_run(p, eng, props, False)

        doc.add_paragraph()

    # Render each conjugation paradigm as a separate 2-column table
    paradigm_idx = 0
    for label, items in blk.sections:
        # items is now a list of dicts {'heb': str, 'eng': str}
        # (or legacy tuple form from object_suffixes)
        if not items:
            continue

        # Object suffixes section: render as inline list
        if label == 'object_suffixes':
            continue  # Skip for now — will be merged into a suffixes row later

        # If this section has a label like 'Who Did It?' or 'To Whom', render it
        if label and label not in ('',):
            hdr_p = doc.add_paragraph()
            _bidi(hdr_p, False); _spacing(hdr_p, before=100, after=30)
            hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _ltr_run(hdr_p, label, sz=12, bold=True, color=NAVY)

        # Filter to dict-form pairs only
        pairs = [it for it in items if isinstance(it, dict)]
        if not pairs:
            continue

        paradigm_idx += 1
        tbl2 = doc.add_table(rows=len(pairs), cols=2)
        tbl2.style = 'Table Grid'

        for row_idx, pair in enumerate(pairs):
            heb = pair.get('heb', '')
            eng = pair.get('eng', '')

            # Hebrew cell (right)
            p_heb = tbl2.rows[row_idx].cells[0].paragraphs[0]
            _bidi(p_heb, True); _spacing(p_heb, before=20, after=20)
            p_heb.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if heb:
                _rtl_run(p_heb, heb, font='David', sz=14, bold=True)

            # English cell (left)
            p_eng = tbl2.rows[row_idx].cells[1].paragraphs[0]
            _bidi(p_eng, False); _spacing(p_eng, before=20, after=20)
            if eng:
                _ltr_run(p_eng, eng, sz=11)

    doc.add_paragraph()


def build_docx(blocks, out_path, style_table=None, hf_events=None):
    """Render document blocks to a Word .docx file.

    style_table: dict from parse_style_table(), maps style_index → props.
                 When provided, each run gets the correct font, size, bold,
                 italic, and underline from the original DavkaWriter document.
    hf_events:   dict {'header': [events], 'footer': [events]} from
                 extract_header_footer_events().  When provided, header/footer
                 content is placed in Word's headers and footers.
    """
    if style_table is None:
        style_table = {}
    if hf_events is None:
        hf_events = {}

    def _props(sty):
        """Return style props for a style code, with sensible defaults."""
        return style_table.get(sty, {})

    # Pre-compute which body blocks are likely titles/centered
    # Heuristic: single-run short block with bold or large font, sandwiched by blanks
    # OR consecutive short large-font blocks (a grid/list of titles)
    def _looks_like_title(blk_idx):
        """A body block that's probably a centered title."""
        if blk_idx >= len(blocks):
            return False
        b = blocks[blk_idx]
        if not (isinstance(b, TextBlock) and b.role == 'body'):
            return False

        # Get the full text and its runs
        full_text = ''.join(t for _, t in b.runs).strip()
        if not full_text or len(full_text) > 80:
            return False
        if len(b.runs) > 3:
            return False

        # Check font size — must be reasonably large (≥16pt)
        max_size = 0
        any_bold = False
        for sty, _ in b.runs:
            props = style_table.get(sty, {})
            if props.get('size_pt', 0) > max_size:
                max_size = props.get('size_pt', 0)
            if props.get('bold', False):
                any_bold = True
        if max_size < 16 and not any_bold:
            return False

        # Structural boundary check: title if at structural boundary
        is_boundary_start = blk_idx == 0 or (
            isinstance(blocks[blk_idx-1], TextBlock)
            and blocks[blk_idx-1].role in ('blank', 'page_break', 'heading', 'section_hdr')
        )

        # OR: the previous block is itself a detected title (grid-of-titles)
        # This catches cases like a poster with many title-like rows
        is_after_title = (
            blk_idx > 0
            and isinstance(blocks[blk_idx-1], TextBlock)
            and blocks[blk_idx-1].role == 'body'
            and max_size >= 20
            and all(style_table.get(s, {}).get('size_pt', 0) >= 18
                    for s, _ in blocks[blk_idx-1].runs)
            and len(''.join(t for _, t in blocks[blk_idx-1].runs).strip()) < 80
            and len(blocks[blk_idx-1].runs) <= 3
        )

        return is_boundary_start or is_after_title

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    for attr in ('left_margin','right_margin','top_margin','bottom_margin'):
        setattr(sec, attr, Inches(0.9))
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── Populate Word header/footer from extracted DWD events ─────────────
    def _render_hf_events(target_paragraph, events):
        """Render header or footer events into the given Word paragraph."""
        target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ev in events:
            if ev['type'] != 'run':
                continue
            try:
                text = decode_run(ev, with_nikud=True, with_trup=True)
            except Exception:
                continue
            if not text.strip():
                continue
            if _is_marker_run(text):
                continue
            sty = ev['style']
            props = style_table.get(sty, {})
            text_bytes = text.encode('utf-8', 'ignore')
            has_heb_script = any(0x0590 <= ord(c) <= 0x05ff for c in text)
            if has_heb_script:
                run_is_heb = True
            elif _is_clean_english(text_bytes):
                run_is_heb = False
            else:
                run_is_heb = props.get('is_hebrew', False) or _is_hebrew_content(text_bytes)
            _styled_run(target_paragraph, text + ' ', props, run_is_heb)

    if hf_events.get('header'):
        hdr = sec.header
        # Clear existing paragraph
        hdr_p = hdr.paragraphs[0]
        hdr_p.clear()
        _render_hf_events(hdr_p, hf_events['header'])

    if hf_events.get('footer'):
        ftr = sec.footer
        ftr_p = ftr.paragraphs[0]
        ftr_p.clear()
        _render_hf_events(ftr_p, hf_events['footer'])

    in_ysk_body = False
    ysk_item_count = 0  # Hebrew letter numbering for YOU SHOULD KNOW body items

    # Hebrew alphabet letters for traditional numbering
    _HEB_LETTERS = ['א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט', 'י',
                    'יא', 'יב', 'יג', 'יד', 'טו', 'טז', 'יז', 'יח', 'יט', 'כ']

    for blk_idx, blk in enumerate(blocks):
        if isinstance(blk, TextBlock):
            role = blk.role

            if role == 'page_break':
                in_ysk_body = False
                pb_p = doc.add_paragraph()
                run = pb_p.add_run()
                br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
                run._r.append(br)

            elif role == 'blank':
                p = doc.add_paragraph(); _spacing(p, before=40, after=40)
                # Note: do NOT reset in_ysk_body — blank paragraphs naturally
                # appear between the section header and items, and between items.

            elif role == 'section_hdr':
                txt = blk.text.strip()
                p = doc.add_paragraph()
                _bidi(p, False); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, before=240, after=120)
                # Match original PDF: plain centered bold underlined text, no decorative box
                run = p.add_run(txt)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.underline = True
                if txt == 'YOU SHOULD KNOW':
                    in_ysk_body = True
                    ysk_item_count = 0
                else:
                    in_ysk_body = False

            elif role == 'heading':
                p = doc.add_paragraph()
                _bidi(p, True); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(p, before=240, after=120)
                # Match original PDF: bold underlined Hebrew (no color, no border)
                run = p.add_run(blk.text.strip())
                run.font.name = 'David'
                run.font.size = Pt(15)
                run.font.bold = True
                run.font.underline = True
                in_ysk_body = False

            elif role == 'mishna':
                p = doc.add_paragraph()
                _bidi(p, True); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _spacing(p, before=60, after=120, line=320)
                if blk.text.strip():
                    _rtl_run(p, blk.text.strip(), font='David', sz=14)
                in_ysk_body = False

            elif role == 'body':
                # Insert a page break before chapter starts (e.g., "פרק יב")
                # for proper visual separation between Bible chapters.
                if is_chapter_start_block(blk):
                    pb_p = doc.add_paragraph()
                    pb_run = pb_p.add_run()
                    pb_br = OxmlElement('w:br')
                    pb_br.set(qn('w:type'), 'page')
                    pb_run._r.append(pb_br)

                p = doc.add_paragraph()

                # Check if this body block is likely a centered title
                is_title = _looks_like_title(blk_idx)

                # Check if this is a Hebrew section/chapter heading.
                # Block-level check ensures we only catch single short Hebrew headings,
                # not mixed-content paragraphs that happen to contain a heading reference.
                is_section_heading = looks_like_section_heading_block(blk)

                # Check if this is a "subtitle pair": short English summary + Hebrew
                # chapter heading.  We render these as TWO centered paragraphs
                # for proper visual hierarchy.
                is_subtitle_pair = (not is_section_heading
                                    and is_subtitle_pair_block(blk))

                # Determine paragraph directionality from majority of runs
                heb_runs = sum(
                    1 for sty, txt in blk.runs
                    if is_heb(sty) or _is_hebrew_content(txt.encode('utf-8', 'ignore'))
                    or style_table.get(sty, {}).get('is_hebrew', False)
                )
                para_is_heb = heb_runs > len(blk.runs) / 2

                # Detect YOU SHOULD KNOW item early so we can override paragraph
                # direction and alignment before any runs are added.
                full_text_pre = ''.join(t for _, t in blk.runs).strip()
                is_ysk_item = (in_ysk_body and not is_title
                              and not is_section_heading
                              and not is_subtitle_pair
                              and len(full_text_pre) >= 8
                              and ysk_item_count < len(_HEB_LETTERS))

                # For YSK items, override direction to LTR (English numbered list
                # layout). For all other body content, use language-based direction.
                if is_ysk_item:
                    _bidi(p, False)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    _bidi(p, para_is_heb)

                if is_section_heading:
                    # Section heading: centered, bold, larger spacing, navy color
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _spacing(p, before=240, after=80)
                    NAVY_HEADING = RGBColor(0x1F, 0x4E, 0x79)
                    for sty, text in blk.runs:
                        text_bytes = text.encode('utf-8', 'ignore')
                        has_heb_script = any(0x0590 <= ord(c) <= 0x05ff for c in text)
                        if has_heb_script:
                            _rtl_run(p, text, font='David', sz=14, bold=True, color=NAVY_HEADING)
                        elif _is_clean_english(text_bytes):
                            _ltr_run(p, text, sz=12, bold=True, color=NAVY_HEADING)
                        else:
                            props = _props(sty)
                            run_is_heb = props.get('is_hebrew', False) or _is_hebrew_content(text_bytes)
                            _styled_run(p, text, props, run_is_heb)
                    continue

                if is_subtitle_pair:
                    # Subtitle pair: render English on top, Hebrew below, both centered
                    NAVY_SUB = RGBColor(0x1F, 0x4E, 0x79)
                    # Identify which run is English vs Hebrew
                    runs_eng = []
                    runs_heb = []
                    for sty, txt in blk.runs:
                        if any(0x0590 <= ord(c) <= 0x05ff for c in txt):
                            runs_heb.append((sty, txt))
                        else:
                            runs_eng.append((sty, txt))
                    # Remove this paragraph (we're going to add 2 new ones)
                    p._element.getparent().remove(p._element)
                    # Render English subtitle first
                    if runs_eng:
                        p_en = doc.add_paragraph()
                        p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _spacing(p_en, before=180, after=10)
                        _bidi(p_en, False)
                        for sty, txt in runs_eng:
                            _ltr_run(p_en, txt, sz=12, bold=True, color=NAVY_SUB)
                    # Render Hebrew heading below
                    if runs_heb:
                        p_he = doc.add_paragraph()
                        p_he.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _spacing(p_he, before=10, after=80)
                        _bidi(p_he, True)
                        for sty, txt in runs_heb:
                            _rtl_run(p_he, txt, font='David', sz=14, bold=True, color=NAVY_SUB)
                    continue

                if is_title:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _spacing(p, before=120, after=60)
                elif is_ysk_item:
                    # Already set LTR + alignment above; just spacing/indent
                    p.paragraph_format.left_indent = Inches(0.4)
                    p.paragraph_format.first_line_indent = Inches(-0.4)
                    _spacing(p, before=20, after=40)
                elif para_is_heb:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _spacing(p, before=30, after=50)
                else:
                    _spacing(p, before=30, after=50)

                # YOU SHOULD KNOW items: prepend Hebrew letter numbering
                # ("א.", "ב.", "ג.", etc.) before any other run content.
                # Add a Unicode Left-to-Right Mark (U+200E) before the number to
                # ensure paragraph direction is LTR even when content starts with Hebrew.
                if is_ysk_item:
                    heb_letter = _HEB_LETTERS[ysk_item_count]
                    # U+200E is the LTR mark; prefix with it to force LTR layout
                    num_run = p.add_run(f'\u200e{heb_letter}.\u200e ')
                    num_run.font.name = 'Times New Roman'
                    num_run.font.size = Pt(13)
                    num_run.font.bold = True
                    # Force the number run to be LTR direction
                    rPr = num_run._r.get_or_add_rPr()
                    rtl_el = OxmlElement('w:rtl')
                    rtl_el.set(qn('w:val'), '0')
                    rPr.append(rtl_el)
                    ysk_item_count += 1

                    # Strip leading control chars (\r, \n, \t) from the first
                    # text run to keep the number adjacent to the content.
                    if blk.runs:
                        first_sty, first_text = blk.runs[0]
                        cleaned = first_text.lstrip('\r\n\t\x0b\x0c')
                        if cleaned != first_text:
                            blk.runs[0] = (first_sty, cleaned)

                # Detect fill-in-the-blank worksheet entries: Hebrew form
                # immediately followed by an English label, no separator,
                # very short (< 25 chars), 2-3 runs.  These appear in pronoun
                # worksheets where Hebrew form and English label are typeset
                # side-by-side in the original.
                full_text = ''.join(t for _, t in blk.runs)
                runs_list = blk.runs
                # Specific pattern: Hebrew run ending without space followed by
                # English run starting without space, all very short
                is_mixed_short = False
                if (len(full_text) < 25 and 2 <= len(runs_list) <= 3):
                    has_heb = any(_is_hebrew_content(t.encode('utf-8','ignore'))
                                   for _, t in runs_list)
                    has_eng = any(all(c < 0x80 for c in t.encode('utf-8','ignore'))
                                   and any(c.isalpha() for c in t)
                                   for _, t in runs_list)
                    if has_heb and has_eng:
                        # Must have a Hebrew→English transition WITHOUT space at boundary
                        for j in range(len(runs_list) - 1):
                            prev_text = runs_list[j][1]
                            next_text = runs_list[j+1][1]
                            if (prev_text and next_text
                                and not prev_text.endswith(' ')
                                and not next_text.startswith(' ')):
                                # Check transition is heb→eng
                                prev_heb = any(0x0590 <= ord(c) <= 0x05ff for c in prev_text)
                                next_eng = (all(c < 0x80 for c in next_text.encode('utf-8','ignore'))
                                            and any(c.isalpha() for c in next_text))
                                if prev_heb and next_eng:
                                    is_mixed_short = True
                                    break

                prev_was_heb = None
                prev_text_ended_with_space = True  # nothing before, treat as space
                for sty, text in blk.runs:
                    props = _props(sty)
                    # Content trumps style metadata: if text is clearly English
                    # (not Hebrew script), mark it as English regardless of what
                    # the style says.  This handles DWD files where CDocStyle
                    # indexing doesn't cleanly map to run style bytes.
                    text_bytes = text.encode('utf-8', 'ignore')
                    has_heb_script = any(0x0590 <= ord(c) <= 0x05ff for c in text)
                    if has_heb_script:
                        run_is_heb = True
                    elif _is_clean_english(text_bytes):
                        run_is_heb = False
                    else:
                        run_is_heb = (
                            is_heb(sty)
                            or props.get('is_hebrew', False)
                            or _is_hebrew_content(text_bytes)
                        )

                    # Insert separator when switching from Hebrew to English in
                    # short fill-in-the-blank worksheet entries
                    if (is_mixed_short
                        and prev_was_heb is True and run_is_heb is False):
                        sep_run = p.add_run('  →  ')
                        sep_run.font.name = 'Calibri'
                        sep_run.font.size = Pt(9)
                        sep_run.font.color.rgb = GRAY
                    # Otherwise just add a space when transitioning between
                    # languages without one already, so words don't run together
                    elif (prev_was_heb is not None
                          and prev_was_heb != run_is_heb
                          and not prev_text_ended_with_space
                          and text and not text[0].isspace()):
                        space_run = p.add_run(' ')

                    _styled_run(p, text, props, run_is_heb)
                    prev_was_heb = run_is_heb
                    prev_text_ended_with_space = bool(text and text[-1].isspace())

        elif isinstance(blk, ImageBlock):
            # Embed image inline; scale to fit 6-inch page width
            MAX_W = 6.0
            w_in = min(MAX_W, blk.width / 100.0)
            p = doc.add_paragraph()
            _bidi(p, False)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before=80, after=80)
            try:
                doc.add_picture(io.BytesIO(blk.raw), width=Inches(w_in))
                # python-docx adds the picture to the last paragraph; move caption
                cap = doc.add_paragraph()
                _bidi(cap, False); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(cap, before=0, after=80)
                _ltr_run(cap,
                         f'Image {blk.index+1}  ({blk.width}×{blk.height} px)',
                         sz=9, color=GRAY)
            except Exception as _e:
                _ltr_run(p, f'[Image {blk.index+1}: could not embed — {_e}]', sz=10, color=GRAY)
            in_ysk_body = False

        elif isinstance(blk, KeyWordBlock):
            if blk.pairs:
                _add_kw_table(doc, blk.pairs)
            in_ysk_body = False

        elif isinstance(blk, ParshaTopicsBlock):
            # Page break before Parsha Topics (major section)
            pb_p = doc.add_paragraph()
            pb_run = pb_p.add_run()
            pb_br = OxmlElement('w:br')
            pb_br.set(qn('w:type'), 'page')
            pb_run._r.append(pb_br)
            _render_parsha_topics(doc, blk, style_table)
            in_ysk_body = False

        elif isinstance(blk, QABlock):
            # Page break before Q&A section
            pb_p = doc.add_paragraph()
            pb_run = pb_p.add_run()
            pb_br = OxmlElement('w:br')
            pb_br.set(qn('w:type'), 'page')
            pb_run._r.append(pb_br)
            _render_qa(doc, blk, style_table)
            in_ysk_body = False

        elif isinstance(blk, PronounBlock):
            # Page break before Pronoun chart
            pb_p = doc.add_paragraph()
            pb_run = pb_p.add_run()
            pb_br = OxmlElement('w:br')
            pb_br.set(qn('w:type'), 'page')
            pb_run._r.append(pb_br)
            _render_pronoun(doc, blk, style_table)
            in_ysk_body = False

    # Font note page
    doc.add_page_break()
    note_hdr = doc.add_paragraph()
    _bidi(note_hdr, False); _spacing(note_hdr, before=0, after=80)
    _ltr_run(note_hdr, 'Font Information', sz=13, bold=True, color=NAVY)
    for line in FONT_NOTE.split('\n'):
        if not line.strip(): continue
        np = doc.add_paragraph()
        _bidi(np, False); _spacing(np, before=20, after=20)
        if line.startswith('  •'):
            np.paragraph_format.left_indent = Inches(0.3)
            parts = line.strip().lstrip('• ').split('→')
            _ltr_run(np, '• ', sz=11)
            if len(parts) == 2:
                _ltr_run(np, parts[0].strip() + '  ', sz=11, bold=True)
                _ltr_run(np, parts[1].strip(), sz=10, color=GRAY)
            else:
                _ltr_run(np, line.strip().lstrip('• '), sz=11)
        else:
            _ltr_run(np, line.strip(), sz=11)

    doc.save(out_path)


# ── Entry point ───────────────────────────────────────────────────────────────
def convert(inp, out=None, with_nikud=True, with_trup=True):
    """Convert a DWD file to DOCX (or ZIP containing DOCX + loose images).

    Returns the path of the output file (.docx or .zip).
    A ZIP is produced when images are found — it contains:
      • converted.docx  (with images embedded inline)
      • images/img_01.jpg, img_02.jpg, …  (original-quality loose copies)
    """
    p = Path(inp)
    if not p.exists():
        sys.exit(f'File not found: {inp}')

    print(f'Reading  {inp}  ({p.stat().st_size:,} bytes)')
    data = p.read_bytes()

    print('Parsing …')
    evts, fmt_name = parse_dwd(data)
    n_runs  = sum(1 for e in evts if e['type'] == 'run')
    n_paras = sum(1 for e in evts if e['type'] == 'para')
    n_imgs  = sum(1 for e in evts if e['type'] == 'image')
    print(f'  Format: {fmt_name}')
    print(f'  {n_runs} runs, {n_paras} para-breaks, {n_imgs} image(s)')

    print('Building document model …')
    blocks = build_model(evts, with_nikud, with_trup)
    pages  = sum(1 for b in blocks if isinstance(b, TextBlock) and b.role == 'page_break')
    kwtbls = sum(1 for b in blocks if isinstance(b, KeyWordBlock))
    img_blocks = [b for b in blocks if isinstance(b, ImageBlock)]
    print(f'  {len(blocks)} blocks  ({pages} page-breaks, {kwtbls} tables, {len(img_blocks)} images)')

    # Determine output path
    stem = p.stem
    if out is None:
        out = str(p.with_suffix('.zip' if img_blocks else '.docx'))
    # Force .zip extension when images present
    out_path = Path(out)
    if img_blocks and out_path.suffix.lower() != '.zip':
        out_path = out_path.with_suffix('.zip')

    docx_path = out_path.with_suffix('.docx') if img_blocks else out_path

    print('Building DOCX …')
    style_table = parse_style_table(data)

    # Normalize style sizes based on usage frequency: heavily-used styles are
    # body text and need to be capped at body sizes (DavkaWriter sometimes
    # stores 76pt for body pesukim which is too large for Word page layout).
    style_table = _normalize_style_sizes(style_table, evts)

    # Extract header/footer events using the run/para signatures for this format
    run_sig, para_sig, _, _ = _detect_format(data)
    hf_events = extract_header_footer_events(data, run_sig, para_sig) if run_sig else {}

    build_docx(blocks, str(docx_path), style_table, hf_events)

    if img_blocks:
        print(f'Packaging ZIP with {len(img_blocks)} image(s) …')
        with zipfile.ZipFile(str(out_path), 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(str(docx_path), f'{stem}.docx')
            for blk in img_blocks:
                ext = blk.fmt.lower()
                img_name = f'images/img_{blk.index+1:02d}.{ext}'
                zf.writestr(img_name, blk.raw)
                print(f'  Added {img_name}  ({blk.width}×{blk.height} px)')
        docx_path.unlink()   # remove the intermediate .docx
        print(f'Saved → {out_path}')
        return str(out_path)

    print(f'Saved → {out_path}')
    return str(out_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__); sys.exit(0)
    with_nikud = '--no-nikud' not in sys.argv
    with_trup  = '--no-trup'  not in sys.argv
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    result = convert(args[0], args[1] if len(args) > 1 else None, with_nikud, with_trup)
    print(f'Output: {result}')
